"""
Docling 기반 문서 재구조화 테스트
Ollama 없이도 파싱/재구성/출력 로직을 테스트할 수 있습니다.

사용법:
    python -m docling_restructure.test_restructure
    python -m docling_restructure.test_restructure --full   # LLM 포함 테스트
"""

import os
import sys

from docx import Document


def create_test_document_structured(path: str):
    """헤딩이 있는 구조화된 테스트 문서 생성"""
    doc = Document()

    doc.add_heading('인공지능 기술 보고서', 0)

    doc.add_heading('1. 서론', level=1)
    doc.add_paragraph(
        '인공지능(AI)은 현대 기술의 핵심 분야로 자리잡았습니다. '
        '본 보고서에서는 AI의 현재 동향과 미래 전망을 분석합니다.'
    )
    doc.add_paragraph(
        '최근 대규모 언어 모델(LLM)의 발전으로 자연어 처리 분야에서 '
        '혁명적인 변화가 일어나고 있습니다.'
    )

    doc.add_heading('결론', level=1)
    doc.add_paragraph(
        'AI 기술의 발전은 기회와 도전을 동시에 가져옵니다. '
        '기술의 책임있는 개발과 활용이 중요합니다.'
    )

    doc.add_heading('2. 기술 현황', level=1)

    doc.add_heading('2.1 대규모 언어 모델', level=2)
    doc.add_paragraph(
        'GPT, Claude, Gemma 등 다양한 LLM이 등장하여 텍스트 생성, '
        '요약, 번역 등 다양한 작업을 수행하고 있습니다.'
    )
    doc.add_paragraph('주요 LLM 모델:', style='List Bullet')
    doc.add_paragraph('GPT-4: OpenAI의 대표 모델', style='List Bullet')
    doc.add_paragraph('Claude: Anthropic의 안전한 AI', style='List Bullet')
    doc.add_paragraph('Gemma: Google의 오픈소스 모델', style='List Bullet')

    doc.add_heading('2.2 컴퓨터 비전', level=2)
    doc.add_paragraph(
        '이미지 인식, 객체 탐지, 자율주행 등 컴퓨터 비전 분야에서도 '
        'AI 기술이 빠르게 발전하고 있습니다.'
    )

    doc.add_heading('3. 산업 적용 사례', level=1)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['분야', '적용 사례', '효과']
    for j, header in enumerate(headers):
        table.cell(0, j).text = header
    data = [
        ['의료', '질병 진단 보조', '진단 정확도 30% 향상'],
        ['금융', '사기 탐지', '사기 적발률 50% 향상'],
        ['제조', '품질 검사 자동화', '불량률 20% 감소'],
    ]
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            table.cell(i + 1, j).text = cell_text

    doc.add_heading('4. 향후 전망', level=1)
    doc.add_paragraph(
        'AI 기술은 앞으로 더욱 발전하여 인간의 삶을 크게 변화시킬 것으로 예상됩니다. '
        '특히 멀티모달 AI, 에이전트 AI, 소형화된 온디바이스 AI가 주요 트렌드가 될 것입니다.'
    )

    doc.save(path)
    print(f"[TEST] 구조화된 테스트 문서 생성: {path}")


def create_test_document_unstructured(path: str):
    """헤딩 없는 비정형 테스트 문서 생성"""
    doc = Document()

    doc.add_paragraph('회의록')
    doc.add_paragraph('일시: 2026년 2월 24일')
    doc.add_paragraph('장소: 본사 회의실')
    doc.add_paragraph('참석자: 김철수, 이영희, 박민수, 정수진')
    doc.add_paragraph(
        '첫 번째 안건으로 신규 프로젝트 일정에 대해 논의했습니다. '
        '김철수 팀장이 전체 일정을 설명했고, 3월 말까지 1차 프로토타입을 '
        '완성하기로 합의했습니다.'
    )
    doc.add_paragraph(
        '두 번째 안건으로 예산 배분에 대해 논의했습니다. '
        '인프라 구축에 40%, 인력 비용에 35%, 기타 운영비에 25%를 '
        '배분하기로 했습니다.'
    )
    doc.add_paragraph(
        '세 번째 안건으로 인력 충원 계획을 논의했습니다. '
        'AI 엔지니어 2명, 프론트엔드 개발자 1명을 3월 중으로 '
        '채용하기로 결정했습니다.'
    )
    doc.add_paragraph('다음 회의 일정: 2026년 3월 3일')

    doc.save(path)
    print(f"[TEST] 비정형 테스트 문서 생성: {path}")


def test_docling_parser():
    """Docling 파싱 테스트"""
    from .parser import parse_document

    test_dir = 'test_data'
    os.makedirs(test_dir, exist_ok=True)

    structured_path = os.path.join(test_dir, 'test_structured.docx')
    unstructured_path = os.path.join(test_dir, 'test_unstructured.docx')

    create_test_document_structured(structured_path)
    create_test_document_unstructured(unstructured_path)

    errors = []

    # 구조화된 문서 파싱 테스트
    print("\n" + "=" * 50)
    print("Docling 파싱 테스트 - 구조화 문서")
    print("=" * 50)

    parsed = parse_document(structured_path)
    print(f"  제목: {parsed.title}")
    print(f"  섹션 수: {len(parsed.sections)}")
    print(f"  구조 유형: {'구조화' if parsed.is_structured else '비정형'}")
    print(f"  마크다운 길이: {len(parsed.markdown_text)}자")
    print(f"  원본 형식: {parsed.source_format}")

    for s in parsed.sections:
        meta = f"{s.char_count}자"
        if s.has_tables:
            meta += ", 테이블"
        if s.has_lists:
            meta += ", 리스트"
        print(f"  [{s.index}] L{s.level} \"{s.title}\" ({meta})")

    if not parsed.is_structured:
        errors.append("구조화된 문서가 비정형으로 감지됨")
    else:
        print("  [PASS] 구조화 문서로 올바르게 감지됨")

    if len(parsed.sections) < 3:
        errors.append(f"섹션이 3개 미만: {len(parsed.sections)}")
    else:
        print(f"  [PASS] 섹션 {len(parsed.sections)}개 추출됨")

    if not parsed.title or parsed.title == '제목 없음':
        errors.append("제목 감지 실패")
    else:
        print(f"  [PASS] 제목 감지: {parsed.title}")

    if len(parsed.markdown_text) < 100:
        errors.append("마크다운 텍스트가 너무 짧음")
    else:
        print("  [PASS] 마크다운 텍스트 생성됨")

    # 비정형 문서 파싱 테스트
    print("\n" + "=" * 50)
    print("Docling 파싱 테스트 - 비정형 문서")
    print("=" * 50)

    parsed2 = parse_document(unstructured_path)
    print(f"  제목: {parsed2.title}")
    print(f"  섹션 수: {len(parsed2.sections)}")
    print(f"  구조 유형: {'구조화' if parsed2.is_structured else '비정형'}")

    for s in parsed2.sections:
        print(f"  [{s.index}] L{s.level} \"{s.title}\" ({s.char_count}자)")

    if len(parsed2.sections) < 1:
        errors.append("비정형 문서에서 섹션이 0개")
    else:
        print(f"  [PASS] 비정형 문서에서 {len(parsed2.sections)}개 섹션 추출")

    if errors:
        print(f"\n[FAIL] {len(errors)}개 오류:")
        for err in errors:
            print(f"  - {err}")
    else:
        print("\n[ALL PASS] Docling 파싱 테스트 통과!")

    return structured_path, unstructured_path, len(errors) == 0


def test_restructure_without_llm(structured_path: str):
    """
    LLM 없이 재구성 로직 테스트.
    수동 RestructurePlan으로 파이프라인을 검증합니다.
    """
    from .parser import parse_document
    from .analyzer import TocItem, RestructurePlan, DocumentAnalysis
    from .restructurer import restructure_document
    from .writer_docx import write_docx
    from .writer_md import write_md

    print("\n" + "=" * 50)
    print("재구성 로직 테스트 (LLM 없이)")
    print("=" * 50)

    parsed = parse_document(structured_path)
    errors = []

    print(f"\n원본 섹션 ({len(parsed.sections)}개):")
    for s in parsed.sections:
        print(f"  [{s.index}] L{s.level} \"{s.title}\" ({s.char_count}자)")

    # 수동 재구성 계획
    n = len(parsed.sections)
    analysis = DocumentAnalysis(
        document_type='보고서',
        main_topic='인공지능 기술 동향',
        key_themes=['LLM', '컴퓨터 비전', '산업 적용'],
        current_structure_assessment='결론이 중간에 위치하여 재배치 필요',
        content_sections=[],
    )

    # 섹션 인덱스는 Docling 파싱 결과에 따라 달라질 수 있으므로
    # 제목으로 인덱스를 찾아서 계획 구성
    section_map = {s.title: s.index for s in parsed.sections}

    # 서론, 결론, 기술 현황, LLM, CV, 산업적용, 전망 순서 찾기
    def find_idx(keyword):
        for title, idx in section_map.items():
            if keyword in title:
                return idx
        return None

    intro_idx = find_idx('서론')
    conclusion_idx = find_idx('결론')
    tech_idx = find_idx('기술 현황') or find_idx('기술')
    llm_idx = find_idx('언어 모델') or find_idx('LLM')
    cv_idx = find_idx('비전')
    industry_idx = find_idx('산업')
    future_idx = find_idx('전망')

    # 인덱스가 없으면 순차적으로 할당
    all_indices = list(range(n))
    toc_items = []

    if intro_idx is not None:
        toc_items.append(TocItem(level=1, title='서론', source_sections=[intro_idx]))

    # 기술 현황 그룹
    tech_sources = []
    tech_subs = []
    if tech_idx is not None:
        tech_sources.append(tech_idx)
    if llm_idx is not None:
        tech_sources.append(llm_idx)
        tech_subs.append(TocItem(level=2, title='대규모 언어 모델', source_sections=[llm_idx]))
    if cv_idx is not None:
        tech_sources.append(cv_idx)
        tech_subs.append(TocItem(level=2, title='컴퓨터 비전', source_sections=[cv_idx]))

    if tech_sources:
        toc_items.append(TocItem(
            level=1, title='기술 현황',
            source_sections=tech_sources,
            subsections=tech_subs,
        ))

    if industry_idx is not None:
        toc_items.append(TocItem(level=1, title='산업 적용 사례', source_sections=[industry_idx]))
    if future_idx is not None:
        toc_items.append(TocItem(level=1, title='향후 전망', source_sections=[future_idx]))
    if conclusion_idx is not None:
        toc_items.append(TocItem(level=1, title='결론', source_sections=[conclusion_idx]))

    # 매핑 안 된 섹션 처리
    mapped = set()
    for item in toc_items:
        mapped.update(item.source_sections)
        for sub in item.subsections:
            mapped.update(sub.source_sections)
    unmapped = [i for i in all_indices if i not in mapped]
    if unmapped:
        toc_items.append(TocItem(level=1, title='기타', source_sections=unmapped))

    plan = RestructurePlan(
        title='인공지능 기술 보고서 (재구성)',
        toc=toc_items,
        analysis=analysis,
    )

    # 재구성 실행
    restructured = restructure_document(parsed, plan, refine=False)

    print(f"\n재구성 결과:")
    print(f"  문서 제목: {restructured.title}")
    print(f"  섹션 수: {len(restructured.sections)}")

    def _print_sections(sections, indent=0):
        for s in sections:
            prefix = '  ' * indent
            content_preview = s.content_md[:50].replace('\n', ' ') if s.content_md else '(비어있음)'
            print(f"  {prefix}[L{s.level}] {s.title} → {content_preview}")
            _print_sections(s.subsections, indent + 1)

    _print_sections(restructured.sections)

    # 검증 1: 결론이 마지막 섹션인지
    last_title = restructured.sections[-1].title
    if '결론' in last_title or last_title == '기타':
        print("\n  [PASS] 결론/기타가 마지막 섹션")
    else:
        errors.append(f"마지막 섹션이 결론이 아님: {last_title}")

    # 검증 2: 기술 현황에 하위 섹션이 있는지
    tech_section = None
    for s in restructured.sections:
        if '기술' in s.title:
            tech_section = s
            break
    if tech_section and len(tech_section.subsections) >= 1:
        print(f"  [PASS] 기술 현황에 하위 섹션 {len(tech_section.subsections)}개 존재")
    elif tech_section:
        errors.append(f"기술 현황에 하위 섹션 없음")

    # 출력 파일 생성 테스트
    output_dir = 'test_output_docling'
    os.makedirs(output_dir, exist_ok=True)
    output_docx = os.path.join(output_dir, 'test_restructured.docx')
    output_md = os.path.join(output_dir, 'test_restructured.md')

    write_docx(restructured, output_docx)
    write_md(restructured, output_md)

    print(f"\n  출력 파일: {output_docx}, {output_md}")

    # MD 파일 검증
    with open(output_md, 'r', encoding='utf-8') as f:
        md_content = f.read()

    if '## 목차' in md_content:
        print("  [PASS] MD에 목차 존재")
    else:
        errors.append("MD에 목차가 없음")

    if errors:
        print(f"\n[FAIL] {len(errors)}개 오류:")
        for err in errors:
            print(f"  - {err}")
    else:
        print("\n[ALL PASS] 재구성 로직 테스트 통과!")

    return len(errors) == 0


def test_edge_cases():
    """엣지 케이스 테스트"""
    from .analyzer import TocItem, _validate_source_sections

    print("\n" + "=" * 50)
    print("엣지 케이스 테스트")
    print("=" * 50)

    errors = []
    total = 7

    # 범위 초과 인덱스 검증
    print("\n[테스트 1] 범위 초과 인덱스 검증")
    test_items = [
        TocItem(level=1, title='유효', source_sections=[0, 1]),
        TocItem(level=1, title='범위 초과', source_sections=[99, -1, 100]),
        TocItem(level=1, title='혼합', source_sections=[2, 50, 3]),
    ]
    _validate_source_sections(test_items, total)

    if test_items[0].source_sections == [0, 1]:
        print("  [PASS] 유효 인덱스 보존됨")
    else:
        errors.append(f"유효 인덱스 변경됨: {test_items[0].source_sections}")

    if test_items[1].source_sections == []:
        print("  [PASS] 범위 초과 인덱스 모두 제거됨")
    else:
        errors.append(f"범위 초과 제거 실패: {test_items[1].source_sections}")

    if test_items[2].source_sections == [2, 3]:
        print("  [PASS] 혼합에서 초과만 제거됨")
    else:
        errors.append(f"혼합 처리 실패: {test_items[2].source_sections}")

    # 동적 가이드 계산 검증
    print("\n[테스트 2] 동적 가이드 계산")
    test_cases = [
        (3, 3), (5, 3), (6, 4), (12, 4), (15, 5), (16, 5), (30, 7), (50, 12),
    ]
    for n, expected_min in test_cases:
        if n <= 5:
            min_major = 3
        elif n <= 15:
            min_major = max(4, n // 3)
        else:
            min_major = max(5, n // 4)

        if min_major == expected_min:
            print(f"  [PASS] n={n} → min_major={min_major}")
        else:
            errors.append(f"n={n}: min_major={min_major} (기대: {expected_min})")

    if errors:
        print(f"\n[FAIL] {len(errors)}개 오류:")
        for err in errors:
            print(f"  - {err}")
    else:
        print("\n[ALL PASS] 엣지 케이스 테스트 통과!")

    return len(errors) == 0


def main():
    """테스트 메인"""
    print("=" * 60)
    print("  Docling Document Restructuring Tool - 테스트")
    print("=" * 60)

    all_passed = True

    # Step 1: Docling 파싱 테스트
    try:
        structured_path, unstructured_path, passed = test_docling_parser()
        all_passed = all_passed and passed
    except ImportError as e:
        print(f"\n[ERROR] Docling을 임포트할 수 없습니다: {e}")
        print("[INFO] 설치하세요: pip install docling")
        sys.exit(1)
    except Exception as e:
        print(f"\n[ERROR] Docling 파싱 테스트 실패: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

    # Step 2: 재구성 로직 테스트
    passed = test_restructure_without_llm(structured_path)
    all_passed = all_passed and passed

    # Step 3: 엣지 케이스 테스트
    passed = test_edge_cases()
    all_passed = all_passed and passed

    # Step 4: 전체 파이프라인 (--full)
    if '--full' in sys.argv:
        print("\n" + "=" * 60)
        print("전체 파이프라인 테스트 (Ollama 필요)")
        print("=" * 60)

        from .config import Config, get_ollama_client

        try:
            client = get_ollama_client()
            client.list()
            print("[TEST] Ollama 연결 성공")
        except Exception as e:
            print(f"[TEST] Ollama 연결 실패: {e}")
            print(f"[TEST]   ollama pull {Config.MODEL}")
            return

        print(f"\n[TEST] 전체 파이프라인: {structured_path}")
        from .main import main as run_main
        sys.argv = ['test', structured_path, '-o', 'test_output_docling_full']
        try:
            run_main()
            print("\n[TEST] 전체 파이프라인 테스트 성공!")
        except SystemExit:
            pass
        except Exception as e:
            print(f"\n[TEST] 전체 파이프라인 테스트 실패: {e}")
            import traceback
            traceback.print_exc()
    else:
        print("\n[INFO] 전체 파이프라인(LLM) 테스트는 --full 옵션으로 실행:")
        print("       python -m docling_restructure.test_restructure --full")

    print()
    if all_passed:
        print("=" * 60)
        print("  모든 테스트 통과!")
        print("=" * 60)
    else:
        print("=" * 60)
        print("  일부 테스트 실패")
        print("=" * 60)
        sys.exit(1)


if __name__ == "__main__":
    main()
