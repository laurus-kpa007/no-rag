"""
.docx 출력 생성 모듈
재구성된 문서를 Word 파일로 출력합니다.
마크다운 콘텐츠를 python-docx 요소로 변환합니다.
"""

import os
import re
from typing import List

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from .restructurer import RestructuredDocument, RestructuredSection


def write_docx(restructured: RestructuredDocument, output_path: str):
    """
    재구성된 문서를 .docx 파일로 출력합니다.

    Args:
        restructured: 재구성된 문서 객체
        output_path: 출력 파일 경로
    """
    print(f"[출력] .docx 파일 생성 중: {output_path}")

    doc = Document()
    _setup_styles(doc)

    # 문서 제목
    title_para = doc.add_heading(restructured.title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 문서 정보
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run(f"문서 유형: {restructured.document_type} | 주제: {restructured.main_topic}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # 페이지 구분선
    doc.add_paragraph('─' * 50)

    # 목차 삽입
    _add_toc(doc, restructured.sections)

    # 페이지 나누기
    doc.add_page_break()

    # 섹션 내용 작성
    for section in restructured.sections:
        _write_section(doc, section)

    # 출력 디렉토리 생성
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)

    doc.save(output_path)
    print(f"[출력] .docx 파일 저장 완료: {output_path}")


def _setup_styles(doc: Document):
    """문서 기본 스타일 설정"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(11)

    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.15

    for i in range(1, 4):
        heading_style_name = f'Heading {i}'
        if heading_style_name in doc.styles:
            hs = doc.styles[heading_style_name]
            hs.font.name = 'Malgun Gothic'
            sizes = {1: 18, 2: 15, 3: 13}
            hs.font.size = Pt(sizes.get(i, 12))
            hs.font.bold = True
            colors = {
                1: RGBColor(0, 51, 102),
                2: RGBColor(0, 76, 153),
                3: RGBColor(51, 102, 153),
            }
            hs.font.color.rgb = colors.get(i, RGBColor(0, 0, 0))


def _add_toc(doc: Document, sections=None):
    """Word 자동 목차(TOC) 필드 삽입 + 수동 목차 미리보기"""
    doc.add_heading('목차', level=1)

    if sections:
        _add_manual_toc_preview(doc, sections)
        doc.add_paragraph()

    # TOC 필드 코드
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar_begin)

    run2 = paragraph.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run2._element.append(instrText)

    run3 = paragraph.add_run()
    fldChar_separate = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._element.append(fldChar_separate)

    run4 = paragraph.add_run('[Word에서 Ctrl+A → F9로 자동 목차를 업데이트하세요]')
    run4.font.color.rgb = RGBColor(128, 128, 128)
    run4.font.size = Pt(9)

    run5 = paragraph.add_run()
    fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._element.append(fldChar_end)


def _add_manual_toc_preview(doc: Document, sections, indent_level=0):
    """수동 목차 미리보기 작성"""
    for section in sections:
        if not section.content_md.strip() and not section.subsections:
            continue
        numbering = _get_section_numbering(section, sections, indent_level)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(indent_level * 1.0)
        run = p.add_run(f"{numbering}{section.title}")
        run.font.size = Pt(10)
        if indent_level == 0:
            run.bold = True
        if section.subsections:
            _add_manual_toc_preview(doc, section.subsections, indent_level + 1)


def _get_section_numbering(section, siblings, indent_level):
    """섹션 번호 생성"""
    visible_index = 0
    for s in siblings:
        if not s.content_md.strip() and not s.subsections:
            continue
        visible_index += 1
        if s is section:
            break
    return f"{visible_index}. " if indent_level == 0 else f"{visible_index}) "


def _write_section(doc: Document, section: RestructuredSection):
    """섹션을 문서에 작성"""
    if not section.content_md.strip() and not section.subsections:
        return

    heading_level = min(max(section.level, 1), 9)
    doc.add_heading(section.title, level=heading_level)

    # 마크다운 콘텐츠를 docx 요소로 변환
    if section.content_md.strip():
        _write_markdown_content(doc, section.content_md)

    for subsection in section.subsections:
        _write_section(doc, subsection)


def _write_markdown_content(doc: Document, md_content: str):
    """마크다운 콘텐츠를 docx 요소로 변환하여 작성"""
    lines = md_content.split('\n')
    i = 0
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 빈 줄
        if not stripped:
            if in_table and table_rows:
                _write_table_from_md(doc, table_rows)
                table_rows = []
                in_table = False
            i += 1
            continue

        # 마크다운 테이블
        if stripped.startswith('|') and '|' in stripped[1:]:
            in_table = True
            # 구분선은 건너뜀
            if not re.match(r'^[|\-\s:]+$', stripped):
                cells = [c.strip() for c in stripped.split('|')[1:-1]]
                if cells:
                    table_rows.append(cells)
            i += 1
            continue

        # 테이블 종료
        if in_table and table_rows:
            _write_table_from_md(doc, table_rows)
            table_rows = []
            in_table = False

        # 리스트 아이템
        list_match = re.match(r'^(\s*)([-*+]|\d+[.)])\s+(.+)$', stripped)
        if list_match:
            content = list_match.group(3)
            doc.add_paragraph(content, style='List Bullet')
            i += 1
            continue

        # 볼드 텍스트 (서브헤딩)
        bold_match = re.match(r'^\*\*(.+)\*\*$', stripped)
        if bold_match:
            p = doc.add_paragraph()
            run = p.add_run(bold_match.group(1))
            run.bold = True
            i += 1
            continue

        # 일반 단락
        doc.add_paragraph(stripped)
        i += 1

    # 마지막 테이블 처리
    if in_table and table_rows:
        _write_table_from_md(doc, table_rows)


def _write_table_from_md(doc: Document, rows: List[List[str]]):
    """마크다운 테이블 데이터를 docx 테이블로 작성"""
    if not rows:
        return

    num_rows = len(rows)
    num_cols = max(len(row) for row in rows) if rows else 0

    if num_cols == 0:
        return

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = row.cells[j]
                cell.text = cell_text
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    doc.add_paragraph()
