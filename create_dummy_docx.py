from docx import Document

doc = Document()
doc.add_heading('Ollama No-RAG Test Document', 0)

doc.add_paragraph('이 문서는 로컬 LLM Q&A 봇 테스트를 위한 예제 문서입니다.')
doc.add_paragraph('주요 내용:')
doc.add_paragraph('1. 벡터 데이터베이스를 사용하지 않습니다.', style='List Bullet')
doc.add_paragraph('2. 문서 전체를 프롬프트에 넣습니다.', style='List Bullet')
doc.add_paragraph('3. Python-docx로 텍스트를 추출합니다.', style='List Bullet')

doc.add_heading('추가 정보', level=1)
doc.add_paragraph('Ollama는 로컬에서 LLM을 실행할 수 있게 해주는 도구입니다.')

table = doc.add_table(rows=2, cols=2)
cell = table.cell(0, 0)
cell.text = '항목'
cell = table.cell(0, 1)
cell.text = '설명'
cell = table.cell(1, 0)
cell.text = 'Context Stuffing'
cell = table.cell(1, 1)
cell.text = '전체 문맥을 한번에 주입하는 방식'

doc.save('data.docx')
print("data.docx created.")
