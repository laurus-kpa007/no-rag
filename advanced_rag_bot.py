# 필요한 라이브러리 설치:
# pip install ollama chromadb rank_bm25 python-docx

import os
import sys
import json
import pickle
import hashlib
import re
import ollama
import chromadb
from rank_bm25 import BM25Okapi
from docx import Document
from chromadb.utils import embedding_functions

# ==========================================
# 1. 설정 (Configuration)
# ==========================================
class Config:
    DOC_PATH = 'data.docx'

    # Ollama 서버 설정 (필요시 수정)
    OLLAMA_HOST = 'http://localhost:11434'

    # 모델 설정
    MODEL_CHAT = 'gemma3:12b'      # 답변 생성용 LLM
    MODEL_EMBED = 'bge-m3' # 임베딩용 모델 (없으면 'ollama pull nomic-embed-text')
    MODEL_RERANK = 'gemma3:12b'    # 리랭킹용 LLM (가벼운 모델 권장)

    # 검색 설정
    CHUNK_SIZE = 500              # 청크 크기 (글자 수)
    CHUNK_OVERLAP = 50            # 청크 겹침 크기
    SEARCH_TOP_K = 5              # 검색시 가져올 문서 수
    NUM_CTX = 32768               # LLM 컨텍스트 윈도우

    # Query Router 설정
    SUMMARY_CHUNK_SIZE = 3000     # 계층적 요약 시 청크 크기
    MAX_CONTEXT_RATIO = 0.7       # 전체 문서 투입 가능 비율 (NUM_CTX 대비)
    PRE_SUMMARIZE = True          # 인덱싱 시 사전 요약 생성 여부

    # 캐시 설정
    CACHE_ENABLED = True          # 임베딩 캐시 사용 여부
    CACHE_DIR = '.rag_cache'      # 캐시 저장 디렉토리

    # PageIndex 설정
    PAGEINDEX_MAX_INSPECT_LOOPS = 3    # 에이전틱 탐색 최대 반복 횟수
    PAGEINDEX_SUMMARY_MAX_CHARS = 200  # 노드 요약 최대 글자 수

    # LLM 하이퍼파라미터 설정
    TEMPERATURE = 0.7             # 응답 창의성 (0.0=결정적, 1.0=창의적, 2.0=매우 랜덤)
    TOP_P = 0.9                   # Nucleus sampling (0.0~1.0, 낮을수록 집중적)
    TOP_K = 40                    # Top-K sampling (높을수록 다양한 토큰 고려)
    REPEAT_PENALTY = 1.1          # 반복 페널티 (1.0=페널티 없음, 높을수록 반복 억제)

# Ollama 클라이언트 (원격 호스트 지원)
def get_ollama_client():
    """OLLAMA_HOST 설정을 사용하는 클라이언트 반환"""
    return ollama.Client(host=Config.OLLAMA_HOST)

def get_llm_options():
    """LLM 호출 시 사용할 options 딕셔너리 반환"""
    return {
        'num_ctx': Config.NUM_CTX,
        'temperature': Config.TEMPERATURE,
        'top_p': Config.TOP_P,
        'top_k': Config.TOP_K,
        'repeat_penalty': Config.REPEAT_PENALTY,
    }

# ==========================================
# 1.5 인덱스 캐시 (Index Cache)
# ==========================================
class IndexCache:
    """
    임베딩, 청크, 요약 데이터를 캐싱하여 재인덱싱을 방지합니다.
    파일 크기, 수정일시, 설정값을 기반으로 캐시 유효성을 검증합니다.
    """
    def __init__(self, file_path):
        self.file_path = os.path.abspath(file_path)
        self.cache_dir = Config.CACHE_DIR
        self.meta_file = os.path.join(self.cache_dir, 'cache_meta.json')
        self.embeddings_file = os.path.join(self.cache_dir, 'embeddings.pkl')
        self.chunks_file = os.path.join(self.cache_dir, 'chunks.pkl')
        self.summary_file = os.path.join(self.cache_dir, 'summary.txt')

        # 캐시 디렉토리 생성
        if Config.CACHE_ENABLED and not os.path.exists(self.cache_dir):
            os.makedirs(self.cache_dir)

    def _get_file_meta(self):
        """현재 파일의 메타데이터를 반환"""
        stat = os.stat(self.file_path)
        return {
            'file_path': self.file_path,
            'file_size': stat.st_size,
            'file_mtime': stat.st_mtime,
            # 설정값도 포함 (설정 변경 시 캐시 무효화)
            'chunk_size': Config.CHUNK_SIZE,
            'chunk_overlap': Config.CHUNK_OVERLAP,
            'model_embed': Config.MODEL_EMBED,
        }

    def _load_cached_meta(self):
        """저장된 캐시 메타데이터를 로드"""
        if not os.path.exists(self.meta_file):
            return None
        try:
            with open(self.meta_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return None

    def is_valid(self):
        """
        캐시가 유효한지 검증합니다.
        파일 크기, 수정일시, 설정값이 모두 일치해야 유효합니다.
        """
        if not Config.CACHE_ENABLED:
            return False

        cached_meta = self._load_cached_meta()
        if not cached_meta:
            return False

        current_meta = self._get_file_meta()

        # 모든 메타데이터가 일치하는지 확인
        for key in current_meta:
            if cached_meta.get(key) != current_meta[key]:
                print(f"   [Cache] 캐시 무효화: {key} 변경됨")
                return False

        # 캐시 파일들이 존재하는지 확인
        if not os.path.exists(self.embeddings_file):
            return False
        if not os.path.exists(self.chunks_file):
            return False

        return True

    def save(self, chunks, embeddings, summary=None):
        """캐시 데이터를 저장"""
        if not Config.CACHE_ENABLED:
            return

        # 메타데이터 저장
        meta = self._get_file_meta()
        meta['cached_at'] = os.path.getmtime(self.file_path)
        with open(self.meta_file, 'w', encoding='utf-8') as f:
            json.dump(meta, f, ensure_ascii=False, indent=2)

        # 청크 저장
        with open(self.chunks_file, 'wb') as f:
            pickle.dump(chunks, f)

        # 임베딩 저장
        with open(self.embeddings_file, 'wb') as f:
            pickle.dump(embeddings, f)

        # 요약 저장 (있는 경우)
        if summary:
            with open(self.summary_file, 'w', encoding='utf-8') as f:
                f.write(summary)

        print(f"   [Cache] 캐시 저장 완료: {self.cache_dir}/")

    def load(self):
        """캐시 데이터를 로드"""
        chunks = None
        embeddings = None
        summary = None

        try:
            with open(self.chunks_file, 'rb') as f:
                chunks = pickle.load(f)
            with open(self.embeddings_file, 'rb') as f:
                embeddings = pickle.load(f)
            if os.path.exists(self.summary_file):
                with open(self.summary_file, 'r', encoding='utf-8') as f:
                    summary = f.read()
        except Exception as e:
            print(f"   [Cache] 캐시 로드 실패: {e}")
            return None, None, None

        return chunks, embeddings, summary

    def clear(self):
        """캐시를 삭제"""
        for f in [self.meta_file, self.embeddings_file, self.chunks_file, self.summary_file]:
            if os.path.exists(f):
                os.remove(f)
        print("   [Cache] 캐시 삭제 완료")

# ==========================================
# 2. 문서 로더 & 청킹 (Loader & Chunking)
# ==========================================
def load_document(file_path):
    if not os.path.exists(file_path):
        return None
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    try:
        if ext == '.docx':
            doc = Document(file_path)
            full_text = []

            # [강력한 추출 방법] 
            # 문서의 구조(단락, 표)를 따르지 않고, 내부 XML에서 모든 텍스트 태그(<w:t>)를 검색합니다.
            # 이 방법은 텍스트 상자, 도형, 표 등 모든 위치의 텍스트를 가져올 수 있습니다.
            
            # 1. Main Body Text
            from docx.oxml.ns import qn
            
            def get_all_text_from_element(element):
                text_list = []
                # w:t (text) 태그를 모두 찾음
                for t in element.findall('.//' + qn('w:t')):
                     if t.text:
                         text_list.append(t.text)
                return text_list

            # 문서 본문(Body) 전체 스캔
            body_texts = get_all_text_from_element(doc.element.body)
            full_text.extend(body_texts)

            text = "\n".join(full_text)
            
            # 너무 붙어서 나오면 가독성이 떨어지므로, 적절히 줄바꿈 보정이 필요할 수 있으나
            # 우선 누락 방지가 최우선이므로 simple join 사용.

        elif ext in ['.md', '.txt']:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        
        # [DEBUG] 추출된 텍스트 저장 (사용자 확인용)
        with open("debug_extracted.txt", "w", encoding="utf-8") as f:
            f.write(text)
        print(f"[System] 추출된 텍스트를 'debug_extracted.txt'에 저장했습니다. 내용이 맞는지 확인해보세요.")
            
    except Exception as e:
        print(f"문서 로드 오류: {e}")
        return None
    return text

def chunk_text(text, size, overlap):
    chunks = []
    start = 0
    while start < len(text):
        end = start + size
        chunk = text[start:end]
        chunks.append(chunk)
        start = end - overlap
    return chunks

# ==========================================
# 3. 검색 엔진 (Retrievers)
# ==========================================
class VectorStore:
    def __init__(self):
        self.client = chromadb.Client() # 메모리 모드 (휘발성)
        # Ollama 임베딩 함수 사용 (chromadb 기본 지원 아님, 커스텀 필요)
        # 여기서는 편의상 SentenceTransformer 대신 간단히 구현하거나
        # Ollama API를 직접 호출해서 임베딩을 가져와야 함.
        # 편의를 위해 chromadb의 DefaultEmbeddingFunction을 쓰지 않고 직접 주입 방식 사용.
        self.collection = self.client.create_collection(name="docs")

    def add_documents(self, chunks, embeddings=None):
        """
        문서를 벡터 DB에 추가합니다.
        embeddings가 제공되면 그대로 사용하고, 없으면 새로 생성합니다.
        """
        ids = [str(i) for i in range(len(chunks))]

        # 임베딩이 제공되지 않은 경우 새로 생성
        if embeddings is None:
            embeddings = []
            print("   [Vector] 임베딩 생성 중... (시간이 걸릴 수 있습니다)")
            client = get_ollama_client()  # 원격 호스트 지원
            for i, chunk in enumerate(chunks):
                try:
                    res = client.embeddings(model=Config.MODEL_EMBED, prompt=chunk)
                    embeddings.append(res['embedding'])
                    if (i + 1) % 10 == 0:
                        print(f"   [Vector] {i+1}/{len(chunks)} 청크 임베딩 완료...")
                except Exception as e:
                    print(f"   [Error] 임베딩 생성 실패: {e}")
                    embeddings.append([0.0]*768) # 더미(실패 시)
        else:
            print(f"   [Vector] 캐시된 임베딩 사용 ({len(embeddings)}개)")

        self.collection.add(
            documents=chunks,
            embeddings=embeddings,
            ids=ids
        )
        print(f"   [Vector] {len(chunks)}개 청크 벡터 DB 저장 완료.")
        return embeddings  # 캐싱을 위해 임베딩 반환

    def search(self, query, top_k=5):
        # 쿼리 임베딩
        client = get_ollama_client()  # 원격 호스트 지원
        res = client.embeddings(model=Config.MODEL_EMBED, prompt=query)
        query_embedding = res['embedding']
        
        results = self.collection.query(
            query_embeddings=[query_embedding],
            n_results=top_k
        )
        # ChromaDB 결과 포맷 정리
        return results['documents'][0] if results['documents'] else []

class KeywordStore:
    def __init__(self):
        self.bm25 = None
        self.chunks = []

    def add_documents(self, chunks):
        self.chunks = chunks
        # 간단한 공백 단위 토크나이징 (한국어 형태소 분석기 없이 약식 구현)
        tokenized_corpus = [doc.split(" ") for doc in chunks]
        self.bm25 = BM25Okapi(tokenized_corpus)
        print(f"   [Keyword] BM25 인덱싱 완료.")

    def search(self, query, top_k=5, expanded_keywords=None):
        """
        BM25 검색을 수행합니다.
        expanded_keywords가 제공되면 해당 키워드로 검색, 없으면 query를 토큰화하여 검색
        """
        if expanded_keywords:
            tokenized_query = expanded_keywords
        else:
            tokenized_query = query.split(" ")

        # 점수 계산
        doc_scores = self.bm25.get_scores(tokenized_query)
        # 상위 k개 인덱스 추출
        top_n_indexes = sorted(range(len(doc_scores)), key=lambda i: doc_scores[i], reverse=True)[:top_k]
        return [self.chunks[i] for i in top_n_indexes]

# ==========================================
# 4. Reranker & Query Refiner (LLM 기반)
# ==========================================
def correct_and_expand_query(query, context_chunks=None):
    """
    사용자 질문의 오타 교정 + 동의어/키워드 확장을 동시에 수행합니다.
    한 번의 LLM 호출로 두 가지를 처리하여 효율성을 높입니다.

    Returns:
        dict: {
            'corrected': 교정된 질문,
            'keywords': 확장된 키워드 리스트 (BM25용)
        }
    """
    context_instruction = ""
    if context_chunks:
        snippets = "\n".join([c[:200] + "..." for c in context_chunks])
        context_instruction = f"""
[참고 문서 내용]
{snippets}

위 문서에 등장하는 전문 용어나 표현을 우선적으로 사용하세요.
"""

    prompt = f"""당신은 검색 최적화 전문가입니다.
아래 [질문]을 분석하여 두 가지 작업을 수행하세요.
{context_instruction}
[질문]
{query}

[작업 1] 오타/띄어쓰기 교정
- 오타와 띄어쓰기를 교정한 문장

[작업 2] 검색 키워드 확장
- 핵심 명사/용어 추출
- 동의어 추가 (예: 제출=마감=신청, 기한=일정=날짜)
- 조사/어미 제거한 원형

다음 형식으로만 출력하세요:
교정: [교정된 질문]
키워드: [키워드1, 키워드2, 키워드3, ...]"""

    try:
        client = get_ollama_client()
        res = client.chat(model=Config.MODEL_CHAT, messages=[{'role': 'user', 'content': prompt}], options=get_llm_options())
        content = res['message']['content'].strip()

        # 응답 파싱
        corrected = query
        keywords = []

        for line in content.split('\n'):
            line = line.strip()
            if line.startswith('교정:'):
                corrected = line.replace('교정:', '').strip().replace('"', '').replace("'", "")
            elif line.startswith('키워드:'):
                kw_str = line.replace('키워드:', '').strip()
                # [키워드1, 키워드2] 또는 키워드1, 키워드2 형식 처리
                kw_str = kw_str.replace('[', '').replace(']', '')
                keywords = [k.strip() for k in kw_str.split(',') if k.strip()]

        # 키워드가 비어있으면 기본 토큰화
        if not keywords:
            keywords = query.split()

        return {
            'corrected': corrected if corrected else query,
            'keywords': keywords
        }
    except:
        return {
            'corrected': query,
            'keywords': query.split()
        }


def rerank_documents(query, docs):
    """
    LLM을 사용하여 문서의 연관성을 평가하고 재정렬합니다.
    """
    print("   [Rerank] 문서 재순위 지정 중(LLM)...")
    scored_docs = []
    client = get_ollama_client()  # 원격 호스트 지원

    for doc in docs:
        prompt = f"""
당신은 검색 품질 관리자입니다.
아래의 [문서]가 [사용자 질문]에 답변하는 데 얼마나 유용한지 평가하세요.

[사용자 질문]
{query}

[문서]
{doc[:500]}...

이 문서가 질문과 관련이 있다면 'Yes', 전혀 관련이 없다면 'No'라고만 답하세요.
설명은 하지 마세요.
"""
        try:
            res = client.chat(model=Config.MODEL_RERANK, messages=[{'role': 'user', 'content': prompt}], options=get_llm_options())
            content = res['message']['content'].strip().lower()
            score = 1 if 'yes' in content else 0
            if score == 1:
                scored_docs.append(doc)
        except:
            pass
            
    return scored_docs if scored_docs else docs[:3] # 실패 시 기본 반환

# ==========================================
# 5. Query Router (질문 유형 분류)
# ==========================================

# 질문 유형 상수
class QueryType:
    SEARCH = "SEARCH"       # 특정 정보 검색
    SUMMARY = "SUMMARY"     # 전체 요약
    COMPARE = "COMPARE"     # 비교/대조
    LIST = "LIST"           # 목록/나열

# 키워드 기반 빠른 분류
QUERY_PATTERNS = {
    QueryType.SUMMARY: [
        # 요약/개요 요청
        "요약", "정리", "개요", "전체", "전반", "대략", "간단히", "핵심",
        "summarize", "summary", "overview", "briefly", "overall", "gist",
        "뭔 내용", "무슨 내용", "어떤 내용", "내용이 뭐", "알려줘 전체",
        # 매뉴얼/가이드 요청 (전체 절차를 원함)
        "매뉴얼", "메뉴얼", "가이드", "안내", "절차", "프로세스", "순서",
        "어떻게 해", "어떻게 하", "방법 알려", "방법을 알려", "전체 과정",
        "manual", "guide", "process", "procedure", "how to", "step by step",
        # 주요 내용 요청
        "주요 내용", "주요내용", "중요한 내용", "핵심 내용", "핵심내용",
        "대강", "대충", "간략", "간략히", "짧게", "요점",
    ],
    QueryType.LIST: [
        # 목록/나열 요청
        "모든", "전부", "목록", "리스트", "나열", "종류", "몇 가지", "몇가지",
        "all", "list", "every", "types", "종류별", "항목",
        # 추가 패턴
        "어떤 것들", "뭐가 있", "뭐뭐", "무엇무엇", "몇 개", "몇개",
        "각각", "하나씩", "전체 목록", "모두 알려", "다 알려",
        "what are", "which ones", "enumerate",
    ],
    QueryType.COMPARE: [
        # 비교/대조 요청
        "비교", "차이", "vs", "versus", "장단점", "다른점", "공통점",
        "compare", "difference", "pros and cons", "versus", "differ",
        # 추가 패턴
        "뭐가 다", "뭐가 달라", "어떻게 다", "차이점", "다른 점",
        "뭐가 좋", "뭐가 나", "어느 게 나", "어떤 게 나",
        "대비", "versus", "and vs", "or vs",
    ],
}

def classify_query_fast(query):
    """
    키워드 기반 빠른 질문 유형 분류 (LLM 호출 없음)
    """
    query_lower = query.lower()

    for query_type, keywords in QUERY_PATTERNS.items():
        if any(kw in query_lower for kw in keywords):
            return query_type

    return QueryType.SEARCH  # 기본값

def classify_query_llm(query):
    """
    LLM 기반 정확한 질문 유형 분류
    """
    prompt = f"""다음 질문의 유형을 분류하세요.

질문: {query}

유형:
- SEARCH: 특정 정보를 찾는 질문 (예: "XG-200 출력은?", "A의 가격은?", "B가 뭐야?")
- SUMMARY: 전체 요약, 개요, 매뉴얼/가이드 요청 (예: "문서 요약해줘", "전체 절차 알려줘", "매뉴얼 알려줘", "어떻게 하는지 알려줘")
- COMPARE: 비교/대조 요청 (예: "A와 B 차이점", "장단점 비교해줘")
- LIST: 목록/나열 요청 (예: "모든 제품 목록", "기능 전부 알려줘", "종류가 뭐가 있어")

유형 하나만 출력하세요 (SEARCH/SUMMARY/COMPARE/LIST):"""

    try:
        client = get_ollama_client()
        res = client.chat(model=Config.MODEL_CHAT, messages=[{'role': 'user', 'content': prompt}], options=get_llm_options())
        result = res['message']['content'].strip().upper()

        # 유효한 유형인지 확인
        if result in [QueryType.SEARCH, QueryType.SUMMARY, QueryType.COMPARE, QueryType.LIST]:
            return result
        # 결과에서 유형 추출 시도
        for qt in [QueryType.SUMMARY, QueryType.LIST, QueryType.COMPARE, QueryType.SEARCH]:
            if qt in result:
                return qt
        return QueryType.SEARCH
    except:
        return classify_query_fast(query)  # 실패 시 키워드 기반으로 폴백

def hierarchical_summary(full_doc, chunk_size=None):
    """
    문서가 너무 길 때 계층적 요약을 수행합니다.
    1단계: 각 청크를 요약
    2단계: 요약들을 합쳐서 최종 컨텍스트 생성
    """
    if chunk_size is None:
        chunk_size = Config.SUMMARY_CHUNK_SIZE

    # 문서를 큰 청크로 분할 (오버랩 없이)
    chunks = chunk_text(full_doc, chunk_size, overlap=0)

    print(f"   [Summary] 계층적 요약 시작: {len(chunks)}개 섹션")

    client = get_ollama_client()
    summaries = []

    for i, chunk in enumerate(chunks):
        print(f"   [Summary] 섹션 {i+1}/{len(chunks)} 요약 중...")
        prompt = f"""다음 내용을 3-5문장으로 핵심만 요약하세요.
중요한 수치, 이름, 용어는 반드시 포함하세요.

내용:
{chunk}

요약:"""

        try:
            res = client.chat(model=Config.MODEL_CHAT, messages=[{'role': 'user', 'content': prompt}], options=get_llm_options())
            summary = res['message']['content'].strip()
            summaries.append(f"[섹션 {i+1}]\n{summary}")
        except Exception as e:
            print(f"   [Warning] 섹션 {i+1} 요약 실패: {e}")
            # 실패 시 원본 청크의 앞부분 사용
            summaries.append(f"[섹션 {i+1}]\n{chunk[:500]}...")

    combined = "\n\n".join(summaries)

    # 결합된 요약이 여전히 너무 크면 재귀적으로 다시 요약
    max_context = int(Config.NUM_CTX * Config.MAX_CONTEXT_RATIO * 4)  # 대략 글자 수 추정
    if len(combined) > max_context:
        print(f"   [Summary] 요약 결과가 여전히 큼. 2차 요약 수행...")
        return hierarchical_summary(combined, chunk_size)

    print(f"   [Summary] 계층적 요약 완료 ({len(full_doc)}자 → {len(combined)}자)")
    return combined

def extract_comparison_entities(query):
    """
    비교 질문에서 비교 대상 엔티티를 추출합니다.
    """
    prompt = f"""다음 질문에서 비교 대상이 되는 항목들을 추출하세요.

질문: {query}

쉼표로 구분하여 항목만 출력하세요 (예: A, B, C):"""

    try:
        client = get_ollama_client()
        res = client.chat(model=Config.MODEL_CHAT, messages=[{'role': 'user', 'content': prompt}], options=get_llm_options())
        entities = [e.strip() for e in res['message']['content'].split(',')]
        return entities if entities else [query]
    except:
        return [query]

class SummaryCache:
    """
    사전 생성된 요약을 저장하는 캐시 클래스
    인덱싱 단계에서 미리 요약을 생성해두고 질문 시 바로 사용
    """
    def __init__(self):
        self.full_summary = None      # 전체 문서 요약
        self.section_summaries = []   # 섹션별 요약 리스트
        self.is_ready = False

    def generate(self, full_doc):
        """
        문서 로드 시 사전 요약을 생성합니다.
        """
        print("\n[Pre-Summary] 사전 요약 생성 중... (인덱싱 단계)")

        max_doc_size = int(Config.NUM_CTX * Config.MAX_CONTEXT_RATIO * 4)

        # 문서가 작으면 전체 문서를 그대로 사용
        if len(full_doc) <= max_doc_size:
            self.full_summary = full_doc
            print(f"[Pre-Summary] 문서가 작음({len(full_doc)}자). 전체 문서를 캐시합니다.")
            self.is_ready = True
            return

        # 문서가 크면 계층적 요약 수행
        print(f"[Pre-Summary] 문서가 큼({len(full_doc)}자). 계층적 요약을 수행합니다...")

        chunk_size = Config.SUMMARY_CHUNK_SIZE
        chunks = chunk_text(full_doc, chunk_size, overlap=0)

        print(f"[Pre-Summary] {len(chunks)}개 섹션으로 분할")

        client = get_ollama_client()
        summaries = []

        for i, chunk in enumerate(chunks):
            print(f"   [Pre-Summary] 섹션 {i+1}/{len(chunks)} 요약 중...")
            prompt = f"""다음 내용을 3-5문장으로 핵심만 요약하세요.
중요한 수치, 이름, 용어는 반드시 포함하세요.

내용:
{chunk}

요약:"""

            try:
                res = client.chat(model=Config.MODEL_CHAT, messages=[{'role': 'user', 'content': prompt}], options=get_llm_options())
                summary = res['message']['content'].strip()
                summaries.append(f"[섹션 {i+1}]\n{summary}")
            except Exception as e:
                print(f"   [Warning] 섹션 {i+1} 요약 실패: {e}")
                summaries.append(f"[섹션 {i+1}]\n{chunk[:500]}...")

        self.section_summaries = summaries
        combined = "\n\n".join(summaries)

        # 결합된 요약이 여전히 크면 2차 요약
        if len(combined) > max_doc_size:
            print(f"[Pre-Summary] 1차 요약 결과가 큼({len(combined)}자). 2차 요약 수행...")
            prompt = f"""다음은 문서의 섹션별 요약입니다. 이를 하나의 통합된 요약으로 정리하세요.
핵심 내용, 주요 절차, 중요한 수치를 포함하세요.

{combined}

통합 요약:"""
            try:
                res = client.chat(model=Config.MODEL_CHAT, messages=[{'role': 'user', 'content': prompt}], options=get_llm_options())
                self.full_summary = res['message']['content'].strip()
            except:
                self.full_summary = combined
        else:
            self.full_summary = combined

        print(f"[Pre-Summary] 사전 요약 완료 ({len(full_doc)}자 → {len(self.full_summary)}자)")
        self.is_ready = True

    def get_summary(self):
        """캐시된 요약을 반환합니다."""
        if self.is_ready and self.full_summary:
            return self.full_summary
        return None

# ==========================================
# 6. PageIndex 모드 (계층적 문서 탐색 에이전트)
# ==========================================

class DocumentNode:
    """
    문서 트리의 노드. 각 노드는 섹션(Heading) 또는 최상위 문서를 나타냅니다.
    [부모 섹션 - 자식 문단] 관계를 가진 트리 구조를 형성합니다.
    """
    def __init__(self, node_id, title, level, content=""):
        self.node_id = node_id          # 고유 ID (예: "001", "001-002")
        self.title = title              # 섹션 제목
        self.level = level              # 헤딩 레벨 (0=루트, 1=Heading1, 2=Heading2, ...)
        self.content = content          # 이 섹션의 직속 본문 (하위 섹션 제외)
        self.children = []              # 자식 노드 리스트
        self.summary = ""               # LLM 생성 요약

    def add_child(self, child_node):
        self.children.append(child_node)

    def get_full_content(self):
        """이 노드와 모든 자식의 본문을 포함한 전체 텍스트 반환"""
        parts = []
        if self.content.strip():
            parts.append(self.content.strip())
        for child in self.children:
            parts.append(f"\n[{child.title}]\n{child.get_full_content()}")
        return "\n".join(parts)

    def to_dict(self):
        """직렬화용 딕셔너리 변환"""
        return {
            'node_id': self.node_id,
            'title': self.title,
            'level': self.level,
            'content': self.content,
            'summary': self.summary,
            'children': [c.to_dict() for c in self.children],
        }

    @staticmethod
    def from_dict(d):
        """딕셔너리에서 DocumentNode 복원"""
        node = DocumentNode(d['node_id'], d['title'], d['level'], d['content'])
        node.summary = d.get('summary', '')
        for cd in d.get('children', []):
            node.add_child(DocumentNode.from_dict(cd))
        return node


def parse_docx_to_tree(file_path):
    """
    python-docx를 사용하여 .docx 문서를 계층적 트리 구조로 파싱합니다.
    Heading 스타일을 기준으로 [부모 섹션 - 자식 문단] 관계를 구성합니다.
    표(Table)도 해당 섹션의 본문으로 포함됩니다.
    """
    doc = Document(file_path)
    root = DocumentNode("000", "문서 전체", level=0)
    node_counter = [0]  # 리스트로 감싸서 클로저에서 수정 가능하게

    def next_id():
        node_counter[0] += 1
        return f"{node_counter[0]:03d}"

    # 문서의 모든 블록 요소(paragraph, table)를 순서대로 순회
    # python-docx의 element.body에서 순서를 보장하기 위해 XML 레벨에서 접근
    from docx.oxml.ns import qn

    body = doc.element.body
    elements = []  # (type, obj) 튜플 리스트

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            # Paragraph
            for para in doc.paragraphs:
                if para._element is child:
                    elements.append(('paragraph', para))
                    break
        elif tag == 'tbl':
            # Table
            for table in doc.tables:
                if table._element is child:
                    elements.append(('table', table))
                    break

    # 스택 기반으로 계층 구조 구축
    # stack[i] = 현재 레벨 i에서의 활성 노드
    stack = {0: root}
    current_parent = root
    current_content_parts = []  # 현재 노드에 쌓이는 본문

    def flush_content():
        """현재 쌓인 본문을 현재 부모 노드에 추가"""
        nonlocal current_content_parts
        if current_content_parts:
            text = "\n".join(current_content_parts)
            current_parent.content += ("\n" + text) if current_parent.content else text
            current_content_parts = []

    def get_heading_level(paragraph):
        """Paragraph의 Heading 레벨을 반환. Heading이 아니면 0 반환."""
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith('Heading'):
            try:
                return int(style_name.replace('Heading', '').strip())
            except ValueError:
                return 0
        # 한국어 스타일 처리
        if '제목' in style_name:
            for ch in style_name:
                if ch.isdigit():
                    return int(ch)
            return 1
        return 0

    def table_to_text(table):
        """Table 객체를 텍스트로 변환"""
        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_text.append(" | ".join(cells))
        return "\n".join(rows_text)

    for elem_type, elem in elements:
        if elem_type == 'paragraph':
            heading_level = get_heading_level(elem)
            text = elem.text.strip()

            if heading_level > 0 and text:
                # 새로운 Heading 발견 → 이전 본문을 flush하고 새 노드 생성
                flush_content()

                new_node = DocumentNode(next_id(), text, heading_level)

                # 적절한 부모를 찾음: 현재 레벨보다 낮은(상위) 레벨의 노드
                parent_level = heading_level - 1
                while parent_level >= 0 and parent_level not in stack:
                    parent_level -= 1

                parent = stack.get(parent_level, root)
                parent.add_child(new_node)

                # 스택 업데이트: 현재 레벨에 새 노드 설정
                stack[heading_level] = new_node
                # 현재 레벨보다 높은(하위) 레벨의 스택 항목 제거
                for lvl in list(stack.keys()):
                    if lvl > heading_level:
                        del stack[lvl]

                current_parent = new_node
            else:
                # 일반 본문 텍스트
                if text:
                    current_content_parts.append(text)

        elif elem_type == 'table':
            table_text = table_to_text(elem)
            if table_text.strip():
                current_content_parts.append(f"[표]\n{table_text}")

    # 마지막 남은 본문 flush
    flush_content()

    return root


def parse_text_to_tree(text):
    """
    일반 텍스트(.md, .txt)를 계층적 트리로 파싱합니다.
    마크다운 헤딩(#, ##, ###)이나 번호 패턴(제1장, 제1조 등)을 기준으로 구조화합니다.
    """
    root = DocumentNode("000", "문서 전체", level=0)
    node_counter = [0]

    def next_id():
        node_counter[0] += 1
        return f"{node_counter[0]:03d}"

    lines = text.split('\n')
    stack = {0: root}
    current_parent = root
    current_content_parts = []

    # 헤딩 패턴 정의
    heading_patterns = [
        (re.compile(r'^(#{1,6})\s+(.+)'), lambda m: (len(m.group(1)), m.group(2).strip())),
        (re.compile(r'^(제\s*\d+\s*장)\s*(.*)'), lambda m: (1, f"{m.group(1)} {m.group(2)}".strip())),
        (re.compile(r'^(제\s*\d+\s*절)\s*(.*)'), lambda m: (2, f"{m.group(1)} {m.group(2)}".strip())),
        (re.compile(r'^(제\s*\d+\s*조)\s*(.*)'), lambda m: (3, f"{m.group(1)} {m.group(2)}".strip())),
    ]

    def detect_heading(line):
        """줄이 헤딩인지 판별. (level, title) 반환, 아니면 None"""
        stripped = line.strip()
        for pattern, extractor in heading_patterns:
            m = pattern.match(stripped)
            if m:
                return extractor(m)
        return None

    def flush():
        nonlocal current_content_parts
        if current_content_parts:
            t = "\n".join(current_content_parts)
            current_parent.content += ("\n" + t) if current_parent.content else t
            current_content_parts = []

    for line in lines:
        heading = detect_heading(line)
        if heading:
            flush()
            level, title = heading
            new_node = DocumentNode(next_id(), title, level)

            parent_level = level - 1
            while parent_level >= 0 and parent_level not in stack:
                parent_level -= 1
            parent = stack.get(parent_level, root)
            parent.add_child(new_node)

            stack[level] = new_node
            for lvl in list(stack.keys()):
                if lvl > level:
                    del stack[lvl]
            current_parent = new_node
        else:
            if line.strip():
                current_content_parts.append(line.strip())

    flush()
    return root


def generate_node_summaries(node, client=None, depth=0):
    """
    트리의 각 노드에 대해 LLM을 사용하여 요약을 생성합니다.
    리프 노드부터 상향식(Bottom-up)으로 요약을 생성합니다.
    """
    if client is None:
        client = get_ollama_client()

    # 먼저 자식 노드들의 요약을 재귀적으로 생성
    for child in node.children:
        generate_node_summaries(child, client, depth + 1)

    # 루트 노드는 특별 처리
    if node.level == 0:
        child_summaries = "\n".join([f"- {c.title}: {c.summary}" for c in node.children if c.summary])
        node.summary = f"이 문서는 {len(node.children)}개의 주요 섹션으로 구성됨. " + child_summaries[:500]
        return

    # 요약할 내용 구성
    content_for_summary = node.content[:1000] if node.content else ""
    if node.children:
        child_info = "\n".join([f"- 하위섹션 [{c.title}]: {c.summary[:100]}" for c in node.children if c.summary])
        content_for_summary += f"\n\n하위 섹션:\n{child_info}"

    if not content_for_summary.strip():
        node.summary = f"[{node.title}] 섹션 (내용 없음)"
        return

    max_chars = Config.PAGEINDEX_SUMMARY_MAX_CHARS
    prompt = f"""다음 섹션의 내용을 {max_chars}자 이내로 요약하세요.
핵심 키워드, 규정/조항 번호, 중요 수치를 반드시 포함하세요.
설명 없이 요약만 출력하세요.

[섹션 제목] {node.title}

[내용]
{content_for_summary}

요약:"""

    try:
        res = client.chat(
            model=Config.MODEL_CHAT,
            messages=[{'role': 'user', 'content': prompt}],
            options=get_llm_options()
        )
        node.summary = res['message']['content'].strip()[:max_chars]
        indent = "  " * depth
        print(f"   {indent}[PageIndex] 노드 요약 완료: [{node.node_id}] {node.title}")
    except Exception as e:
        node.summary = content_for_summary[:max_chars]
        print(f"   [PageIndex] 요약 실패 ({node.title}): {e}")


def build_table_of_contents(node, depth=0):
    """
    트리를 LLM이 읽을 수 있는 목차(Table of Contents) 형태로 변환합니다.
    각 항목에 노드 ID와 요약을 포함합니다.
    """
    lines = []
    indent = "  " * depth

    if node.level > 0:
        summary_text = node.summary if node.summary else "(요약 없음)"
        lines.append(f"{indent}[{node.node_id}] {node.title} — {summary_text}")

    for child in node.children:
        lines.extend(build_table_of_contents(child, depth + 1))

    return lines


def find_node_by_id(node, target_id):
    """노드 ID로 트리에서 특정 노드를 찾습니다."""
    if node.node_id == target_id:
        return node
    for child in node.children:
        result = find_node_by_id(child, target_id)
        if result:
            return result
    return None


def find_nodes_by_ids(root, id_list):
    """여러 노드 ID로 노드들을 찾아 반환합니다."""
    nodes = []
    for nid in id_list:
        nid = nid.strip()
        found = find_node_by_id(root, nid)
        if found:
            nodes.append(found)
    return nodes


def pageindex_planning(query, toc_text, client=None):
    """
    1단계 (Planning): 전체 목차를 보고 질문에 답하기 위해 탐색할 섹션을 결정합니다.
    LLM이 '감사관(Auditor)'의 페르소나로 목차를 훑어보고 관련 섹션 ID를 반환합니다.
    """
    if client is None:
        client = get_ollama_client()

    prompt = f"""당신은 사내 규정집의 목차를 보고 필요한 조항을 스스로 찾아가는 숙련된 감사관(Auditor)입니다.
벡터 유사도가 아닌, 문서의 논리적 위치를 근거로 판단합니다.

아래는 규정 문서의 전체 목차입니다. 각 항목은 [노드ID] 제목 — 요약 형식입니다.

[문서 목차]
{toc_text}

[사용자 질문]
{query}

위 질문에 답변하기 위해 반드시 확인해야 할 섹션의 노드ID를 선택하세요.
관련성이 높은 순서로 최대 5개까지 선택하세요.

다음 형식으로만 출력하세요:
선택: [노드ID1, 노드ID2, 노드ID3]
이유: (각 섹션을 선택한 이유를 한 줄로 설명)"""

    try:
        res = client.chat(
            model=Config.MODEL_CHAT,
            messages=[{'role': 'user', 'content': prompt}],
            options=get_llm_options()
        )
        content = res['message']['content'].strip()

        # 선택된 노드 ID 파싱
        selected_ids = []
        reason = ""
        for line in content.split('\n'):
            line = line.strip()
            if line.startswith('선택:') or line.startswith('선택 :'):
                ids_str = line.split(':', 1)[1].strip()
                ids_str = ids_str.replace('[', '').replace(']', '')
                selected_ids = [s.strip() for s in ids_str.split(',') if s.strip()]
            elif line.startswith('이유:') or line.startswith('이유 :'):
                reason = line.split(':', 1)[1].strip()

        # ID가 파싱되지 않으면 본문에서 3자리 숫자 패턴 추출
        if not selected_ids:
            selected_ids = re.findall(r'\b(\d{3})\b', content)

        return selected_ids[:5], reason, content
    except Exception as e:
        print(f"   [PageIndex] Planning 실패: {e}")
        return [], "", ""


def pageindex_inspection(query, section_content, section_title, visited_sections, toc_text, client=None):
    """
    2단계 (Inspection): 선택된 섹션의 내용을 읽고 답변 충분성을 판단합니다.
    부족하면 다음에 탐색할 섹션을 제안합니다.
    """
    if client is None:
        client = get_ollama_client()

    visited_list = ", ".join(visited_sections) if visited_sections else "없음"

    prompt = f"""당신은 사내 규정집을 조사하는 감사관(Auditor)입니다.
현재 특정 섹션의 내용을 읽고 있습니다. 사용자의 질문에 충분히 답변할 수 있는지 판단하세요.

[사용자 질문]
{query}

[현재 확인 중인 섹션: {section_title}]
{section_content[:3000]}

[이미 확인한 섹션들]
{visited_list}

판단 결과를 다음 형식으로 출력하세요:
충분: Yes 또는 No
근거문장: (답변의 근거가 되는 핵심 문장을 원문에서 인용)
추가탐색: (No인 경우, 아래 목차에서 추가로 확인할 노드ID. Yes인 경우 "없음")

[참고 목차]
{toc_text[:2000]}"""

    try:
        res = client.chat(
            model=Config.MODEL_CHAT,
            messages=[{'role': 'user', 'content': prompt}],
            options=get_llm_options()
        )
        content = res['message']['content'].strip()

        sufficient = False
        evidence = ""
        next_ids = []

        for line in content.split('\n'):
            line = line.strip()
            if line.startswith('충분:') or line.startswith('충분 :'):
                val = line.split(':', 1)[1].strip().lower()
                sufficient = 'yes' in val
            elif line.startswith('근거문장:') or line.startswith('근거문장 :'):
                evidence = line.split(':', 1)[1].strip()
            elif line.startswith('추가탐색:') or line.startswith('추가탐색 :'):
                ids_str = line.split(':', 1)[1].strip()
                if ids_str != '없음' and ids_str:
                    ids_str = ids_str.replace('[', '').replace(']', '')
                    next_ids = [s.strip() for s in ids_str.split(',') if s.strip()]

        if not next_ids and not sufficient:
            next_ids = re.findall(r'\b(\d{3})\b', content)

        return sufficient, evidence, next_ids, content
    except Exception as e:
        print(f"   [PageIndex] Inspection 실패: {e}")
        return True, "", [], ""


def pageindex_generate_answer(query, collected_sections, client=None):
    """
    수집된 섹션 정보를 기반으로 최종 답변을 생성합니다.
    모든 답변에 [섹션 제목]과 [근거 문장]을 포함합니다.
    """
    if client is None:
        client = get_ollama_client()

    # 수집된 섹션들을 컨텍스트로 구성
    context_parts = []
    for section in collected_sections:
        context_parts.append(
            f"[섹션: {section['title']}] (노드ID: {section['node_id']})\n"
            f"{section['content']}\n"
            f"근거: {section.get('evidence', '해당 섹션 전체 참조')}"
        )

    context = "\n\n---\n\n".join(context_parts)

    prompt = f"""당신은 사내 규정을 정확히 해석하는 감사관(Auditor)입니다.
아래의 [참조 섹션]들을 근거로 사용자의 질문에 답변하세요.

중요 규칙:
1. 반드시 답변의 근거가 되는 **[섹션 제목]**을 명시하세요.
2. 핵심 내용은 원문의 **[근거 문장]**을 직접 인용하세요.
3. 참조 섹션에 없는 내용은 답변하지 마세요.
4. 벡터 유사도가 아닌, 문서의 논리적 위치(목차 구조)를 통해 찾아낸 정보임을 전제하세요.

[참조 섹션]
{context}

[질문]
{query}"""

    return prompt


class PageIndexStore:
    """
    PageIndex 모드의 핵심 저장소.
    문서를 계층적 트리로 파싱하고, 노드 요약을 생성하며,
    에이전틱 추론 탐색을 수행합니다.
    """
    def __init__(self):
        self.root = None              # DocumentNode 트리 루트
        self.toc_text = ""            # LLM용 목차 텍스트
        self.is_ready = False

    def build_index(self, file_path, full_text=None):
        """
        문서를 파싱하여 트리 구조를 구축하고 노드 요약을 생성합니다.
        """
        print("\n[PageIndex] 계층적 문서 인덱싱 시작...")

        ext = os.path.splitext(file_path)[1].lower()

        # 1단계: 문서 → 트리 구조 파싱
        print("   [PageIndex] 1단계: 문서를 트리 구조로 파싱 중...")
        if ext == '.docx':
            self.root = parse_docx_to_tree(file_path)
        else:
            # .md, .txt 등은 텍스트 기반 파싱
            if full_text is None:
                with open(file_path, 'r', encoding='utf-8') as f:
                    full_text = f.read()
            self.root = parse_text_to_tree(full_text)

        # 트리 통계 출력
        total_nodes = self._count_nodes(self.root)
        max_depth = self._max_depth(self.root)
        print(f"   [PageIndex] 트리 구축 완료: {total_nodes}개 노드, 최대 깊이 {max_depth}")

        # 노드가 하나도 없으면 (헤딩이 없는 문서) 폴백
        if total_nodes <= 1 and full_text:
            print("   [PageIndex] 헤딩이 감지되지 않음. 텍스트를 청크 단위로 분할합니다...")
            self._fallback_chunking(full_text)
            total_nodes = self._count_nodes(self.root)
            print(f"   [PageIndex] 폴백 분할 완료: {total_nodes}개 노드")

        # 2단계: 각 노드 요약 생성
        print("   [PageIndex] 2단계: 노드 요약 생성 중... (LLM 호출)")
        generate_node_summaries(self.root)

        # 3단계: 목차 생성
        print("   [PageIndex] 3단계: 목차 구성 중...")
        toc_lines = build_table_of_contents(self.root)
        self.toc_text = "\n".join(toc_lines)
        print(f"   [PageIndex] 목차 생성 완료 ({len(toc_lines)}개 항목)")

        self.is_ready = True
        print("[PageIndex] 인덱싱 완료!\n")

    def _fallback_chunking(self, text):
        """헤딩이 없는 문서를 위한 폴백: 일정 크기로 분할하여 가상 섹션 생성"""
        chunk_size = 2000
        chunks = chunk_text(text, chunk_size, overlap=100)
        self.root = DocumentNode("000", "문서 전체", level=0)
        for i, chunk in enumerate(chunks):
            node_id = f"{i+1:03d}"
            title = f"섹션 {i+1}"
            # 첫 줄에서 제목 추출 시도
            first_line = chunk.strip().split('\n')[0][:50]
            if first_line:
                title = f"섹션 {i+1}: {first_line}"
            node = DocumentNode(node_id, title, level=1, content=chunk)
            self.root.add_child(node)

    def _count_nodes(self, node):
        """트리의 전체 노드 수를 세기"""
        count = 1
        for child in node.children:
            count += self._count_nodes(child)
        return count

    def _max_depth(self, node, current=0):
        """트리의 최대 깊이"""
        if not node.children:
            return current
        return max(self._max_depth(c, current + 1) for c in node.children)

    def search(self, query):
        """
        에이전틱 추론 탐색을 수행합니다.
        1단계: Planning - 목차를 보고 탐색할 섹션 결정
        2단계: Inspection - 섹션 내용을 읽고 충분한지 판단, 부족하면 Loop
        """
        if not self.is_ready:
            return []

        client = get_ollama_client()
        collected_sections = []
        visited_ids = set()

        # 1단계: Planning
        print("   [PageIndex] 1단계(Planning): 목차에서 관련 섹션 탐색 중...")
        selected_ids, reason, _ = pageindex_planning(query, self.toc_text, client)
        print(f"   [PageIndex] 선택된 섹션: {selected_ids}")
        if reason:
            print(f"   [PageIndex] 선택 이유: {reason}")

        if not selected_ids:
            print("   [PageIndex] 관련 섹션을 찾지 못함. 루트 노드의 자식들을 사용합니다.")
            selected_ids = [c.node_id for c in self.root.children[:3]]

        # 2단계: Inspection Loop
        max_loops = Config.PAGEINDEX_MAX_INSPECT_LOOPS
        loop_count = 0
        pending_ids = list(selected_ids)

        while pending_ids and loop_count < max_loops:
            loop_count += 1
            current_id = pending_ids.pop(0)

            if current_id in visited_ids:
                continue
            visited_ids.add(current_id)

            node = find_node_by_id(self.root, current_id)
            if not node:
                print(f"   [PageIndex] 노드 {current_id}을 찾을 수 없음. 건너뜁니다.")
                continue

            section_content = node.get_full_content()
            print(f"   [PageIndex] 2단계(Inspection #{loop_count}): [{node.node_id}] {node.title} 분석 중...")

            sufficient, evidence, next_ids, _ = pageindex_inspection(
                query, section_content, node.title,
                [f"[{nid}]" for nid in visited_ids],
                self.toc_text, client
            )

            collected_sections.append({
                'node_id': node.node_id,
                'title': node.title,
                'content': section_content[:3000],
                'evidence': evidence,
            })

            if sufficient:
                print(f"   [PageIndex] 충분한 정보 확보 완료!")
                break
            else:
                # 아직 확인하지 않은 섹션만 추가
                new_ids = [nid for nid in next_ids if nid not in visited_ids]
                if new_ids:
                    print(f"   [PageIndex] 추가 탐색 필요: {new_ids}")
                    pending_ids = new_ids + pending_ids

        print(f"   [PageIndex] 총 {len(collected_sections)}개 섹션 수집 완료 (탐색 {loop_count}회)")
        return collected_sections

    def get_toc(self):
        """목차 텍스트 반환"""
        return self.toc_text

    def save_to_cache(self, cache_dir):
        """PageIndex 데이터를 캐시에 저장"""
        if not os.path.exists(cache_dir):
            os.makedirs(cache_dir)

        cache_file = os.path.join(cache_dir, 'pageindex.pkl')
        data = {
            'root': self.root.to_dict() if self.root else None,
            'toc_text': self.toc_text,
            'is_ready': self.is_ready,
        }
        with open(cache_file, 'wb') as f:
            pickle.dump(data, f)
        print(f"   [PageIndex Cache] 저장 완료: {cache_file}")

    def load_from_cache(self, cache_dir):
        """PageIndex 데이터를 캐시에서 로드"""
        cache_file = os.path.join(cache_dir, 'pageindex.pkl')
        if not os.path.exists(cache_file):
            return False

        try:
            with open(cache_file, 'rb') as f:
                data = pickle.load(f)
            if data.get('root'):
                self.root = DocumentNode.from_dict(data['root'])
            self.toc_text = data.get('toc_text', '')
            self.is_ready = data.get('is_ready', False)
            print(f"   [PageIndex Cache] 로드 완료")
            return self.is_ready
        except Exception as e:
            print(f"   [PageIndex Cache] 로드 실패: {e}")
            return False


# ==========================================
# 7. 메인 로직
# ==========================================
def main():
    print("=== Advanced Multi-Mode RAG Bot ===")
    
    # 1. 문서 로드
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        print(f"파일 경로(인자): {file_path}")
    else:
        file_path = input(f"분석할 파일 경로 입력 (기본값: {Config.DOC_PATH}): ").strip() or Config.DOC_PATH
    
    text = load_document(file_path)
    if not text: return

    print(f"\n[System] 문서 로드 성공 ({len(text)}자).")

    # 2. 캐시 확인 및 인덱싱
    index_cache = IndexCache(file_path)
    vector_store = VectorStore()
    keyword_store = KeywordStore()
    summary_cache = SummaryCache()
    pageindex_store = PageIndexStore()

    if index_cache.is_valid():
        # 캐시가 유효하면 로드
        print("[Cache] ⚡ 유효한 캐시 발견! 캐시에서 로드합니다...")
        cached_chunks, cached_embeddings, cached_summary = index_cache.load()

        if cached_chunks and cached_embeddings:
            chunks = cached_chunks
            print(f"[Cache] {len(chunks)}개 청크 로드 완료")

            # 벡터 스토어에 캐시된 임베딩 사용
            vector_store.add_documents(chunks, embeddings=cached_embeddings)

            # 키워드 스토어 인덱싱 (BM25는 빠르므로 항상 새로 생성)
            keyword_store.add_documents(chunks)

            # 캐시된 요약 사용
            if cached_summary:
                summary_cache.full_summary = cached_summary
                summary_cache.is_ready = True
                print(f"[Cache] 사전 요약 로드 완료 ({len(cached_summary)}자)")
            elif Config.PRE_SUMMARIZE:
                summary_cache.generate(text)

            # PageIndex 캐시 로드 시도
            if not pageindex_store.load_from_cache(Config.CACHE_DIR):
                print("[PageIndex] 캐시 없음. 질문 시 자동 빌드됩니다.")
        else:
            print("[Cache] 캐시 로드 실패. 새로 인덱싱합니다...")
            index_cache.clear()
            # 아래 else 블록과 동일한 처리
            chunks = chunk_text(text, Config.CHUNK_SIZE, Config.CHUNK_OVERLAP)
            print(f"[System] 총 {len(chunks)}개의 청크 생성됨.")
            embeddings = vector_store.add_documents(chunks)
            keyword_store.add_documents(chunks)
            if Config.PRE_SUMMARIZE:
                summary_cache.generate(text)
            index_cache.save(chunks, embeddings, summary_cache.get_summary())
            print("[PageIndex] 캐시 없음. 질문 시 자동 빌드됩니다.")
    else:
        # 캐시가 없거나 무효화됨 → 새로 인덱싱
        print("[Cache] 캐시 없음 또는 무효화됨. 새로 인덱싱합니다...")
        chunks = chunk_text(text, Config.CHUNK_SIZE, Config.CHUNK_OVERLAP)
        print(f"[System] 총 {len(chunks)}개의 청크 생성됨.")

        # 벡터 스토어 인덱싱 (임베딩 생성)
        embeddings = vector_store.add_documents(chunks)

        # 키워드 스토어 인덱싱
        keyword_store.add_documents(chunks)

        # 사전 요약 생성
        if Config.PRE_SUMMARIZE:
            summary_cache.generate(text)
        else:
            print("[Pre-Summary] 사전 요약 비활성화됨 (Config.PRE_SUMMARIZE=False)")

        # 캐시 저장
        index_cache.save(chunks, embeddings, summary_cache.get_summary())
        print("[PageIndex] 질문 시 자동 빌드됩니다.")

    while True:
        print("\n" + "="*40)
        print(" 1. 파일 전체 검색 (No-RAG, Context Stuffing)")
        print(" 2. 벡터 검색 (Vector Store)")
        print(" 3. 키워드 검색 (BM25)")
        print(" 4. 하이브리드 검색 (Hybrid + Rerank)")
        print(" 5. 자동 모드 (Query Router) ★ 추천")
        print(" 6. PageIndex (계층적 목차 탐색) ★ 규정 문서 특화")
        print(" q. 종료")
        print("="*40)

        mode = input("검색 모드를 선택하세요 (1-6): ").strip()
        if mode.lower() in ['q', 'exit']: break
        
        original_query = input("\n질문: ").strip()
        if not original_query: continue

        # 문맥 기반 질문 교정 + 키워드 확장 (한 번의 LLM 호출)
        print("   [Query] 오타 교정 및 키워드 확장 중...")
        try:
            # 1. 벡터 검색으로 관련 문맥(청크)를 먼저 가져옴 (오타에 강함)
            pre_search_docs = vector_store.search(original_query, top_k=3)
        except Exception as e:
            print(f"   [Warning] 문맥 파악 실패: {e}")
            pre_search_docs = []

        # 오타 교정 + 동의어 확장 동시 수행
        query_result = correct_and_expand_query(original_query, context_chunks=pre_search_docs)
        query = query_result['corrected']
        expanded_keywords = query_result['keywords']

        if query != original_query:
            print(f"   => 교정된 질문: {query}")
        else:
            print(f"   => 질문: {query}")
        print(f"   => 확장 키워드: {expanded_keywords}")

        context = ""
        
        # --- 검색 단계 ---
        if mode == '1':
            context = text # 전체 텍스트
            print(f"[Mode 1] 전체 문서({len(text)}자)를 컨텍스트로 사용합니다.")
            
        elif mode == '2':
            docs = vector_store.search(query, top_k=Config.SEARCH_TOP_K)
            context = "\n---\n".join(docs)
            print(f"[Mode 2] 벡터 유사도 상위 {len(docs)}개 청크 사용.")
            
        elif mode == '3':
            docs = keyword_store.search(query, top_k=Config.SEARCH_TOP_K, expanded_keywords=expanded_keywords)
            context = "\n---\n".join(docs)
            print(f"[Mode 3] 키워드 매칭 상위 {len(docs)}개 청크 사용.")
            
        elif mode == '4':
            vec_docs = vector_store.search(query, top_k=Config.SEARCH_TOP_K)
            key_docs = keyword_store.search(query, top_k=Config.SEARCH_TOP_K, expanded_keywords=expanded_keywords)

            # 중복 제거해서 합치기
            combined_docs = list(set(vec_docs + key_docs))
            print(f"[Hybrid] 1차 검색 완료 ({len(combined_docs)}개 문서). Reranking 시작...")

            # 리랭크
            final_docs = rerank_documents(query, combined_docs)
            context = "\n---\n".join(final_docs)
            print(f"[Mode 4] 최종 선정된 {len(final_docs)}개 청크 사용.")

        elif mode == '5':
            # === 자동 모드 (Query Router) ===
            print("   [Router] 질문 유형 분석 중...")

            # 1. 질문 유형 분류 (빠른 키워드 분류 먼저, 불확실하면 LLM)
            query_type = classify_query_fast(query)

            # 키워드로 SEARCH로 분류되었지만, LLM으로 재확인 (선택적)
            if query_type == QueryType.SEARCH:
                # 더 정확한 분류가 필요하면 LLM 사용 (비용 vs 정확도 트레이드오프)
                query_type = classify_query_llm(query)

            print(f"   [Router] 질문 유형: {query_type}")

            # 2. 유형별 처리
            if query_type == QueryType.SUMMARY:
                # 요약형: 사전 생성된 요약 사용 (없으면 실시간 생성)
                cached = summary_cache.get_summary()
                if cached:
                    context = cached
                    print(f"[Mode 5-SUMMARY] 사전 생성된 요약 사용 ({len(cached)}자) ⚡ 빠른 응답")
                else:
                    # 사전 요약이 없으면 기존 방식으로 실시간 생성
                    max_doc_size = int(Config.NUM_CTX * Config.MAX_CONTEXT_RATIO * 4)
                    if len(text) <= max_doc_size:
                        context = text
                        print(f"[Mode 5-SUMMARY] 전체 문서({len(text)}자)를 컨텍스트로 사용합니다.")
                    else:
                        print(f"[Mode 5-SUMMARY] 문서가 큼({len(text)}자). 계층적 요약을 수행합니다.")
                        context = hierarchical_summary(text)

            elif query_type == QueryType.LIST:
                # 목록형: 더 많은 청크 검색
                extended_top_k = min(Config.SEARCH_TOP_K * 4, len(chunks))
                vec_docs = vector_store.search(query, top_k=extended_top_k)
                key_docs = keyword_store.search(query, top_k=extended_top_k, expanded_keywords=expanded_keywords)
                combined_docs = list(set(vec_docs + key_docs))
                context = "\n---\n".join(combined_docs)
                print(f"[Mode 5-LIST] 확장 검색으로 {len(combined_docs)}개 청크 사용.")

            elif query_type == QueryType.COMPARE:
                # 비교형: 각 엔티티별로 검색 후 병합
                entities = extract_comparison_entities(query)
                print(f"   [Router] 비교 대상: {entities}")

                all_docs = []
                for entity in entities:
                    vec_docs = vector_store.search(entity, top_k=Config.SEARCH_TOP_K)
                    # 엔티티 검색은 해당 엔티티를 키워드로 사용
                    key_docs = keyword_store.search(entity, top_k=Config.SEARCH_TOP_K, expanded_keywords=[entity])
                    all_docs.extend(vec_docs + key_docs)

                combined_docs = list(set(all_docs))
                final_docs = rerank_documents(query, combined_docs)
                context = "\n---\n".join(final_docs)
                print(f"[Mode 5-COMPARE] {len(entities)}개 항목 검색, {len(final_docs)}개 청크 사용.")

            else:  # QueryType.SEARCH
                # 검색형: 하이브리드 검색 + 리랭킹 (기존 모드 4와 동일)
                vec_docs = vector_store.search(query, top_k=Config.SEARCH_TOP_K)
                key_docs = keyword_store.search(query, top_k=Config.SEARCH_TOP_K, expanded_keywords=expanded_keywords)
                combined_docs = list(set(vec_docs + key_docs))
                final_docs = rerank_documents(query, combined_docs)
                context = "\n---\n".join(final_docs)
                print(f"[Mode 5-SEARCH] 하이브리드 검색, {len(final_docs)}개 청크 사용.")

        elif mode == '6':
            # === PageIndex 모드 (계층적 목차 탐색 에이전트) ===
            print("   [PageIndex] 계층적 목차 기반 에이전틱 탐색 모드")

            # PageIndex가 아직 빌드되지 않았으면 빌드
            if not pageindex_store.is_ready:
                print("   [PageIndex] 최초 실행: 문서 인덱스를 구축합니다...")
                pageindex_store.build_index(file_path, full_text=text)
                # 캐시에 저장
                pageindex_store.save_to_cache(Config.CACHE_DIR)

            # 목차 표시 (선택적)
            toc = pageindex_store.get_toc()
            toc_lines = toc.split('\n')
            print(f"\n   [PageIndex] 문서 목차 ({len(toc_lines)}개 섹션):")
            for line in toc_lines[:15]:
                print(f"     {line}")
            if len(toc_lines) > 15:
                print(f"     ... 외 {len(toc_lines) - 15}개 섹션")

            # 에이전틱 탐색 수행
            print(f"\n   [PageIndex] 에이전틱 추론 탐색 시작...")
            collected_sections = pageindex_store.search(query)

            if collected_sections:
                # PageIndex 전용 답변 생성 프롬프트
                prompt_text = pageindex_generate_answer(query, collected_sections)

                print(f"\n[Mode 6] PageIndex 탐색 완료. {len(collected_sections)}개 섹션 기반 답변 생성 중...\n")

                try:
                    client = ollama.Client(host=Config.OLLAMA_HOST)
                    stream = client.chat(
                        model=Config.MODEL_CHAT,
                        messages=[{'role': 'user', 'content': prompt_text}],
                        stream=True,
                        options=get_llm_options()
                    )
                    for chunk_data in stream:
                        print(chunk_data['message']['content'], end="", flush=True)
                    print()
                except Exception as e:
                    print(f"Ollama 오류: {e}")
                continue  # Mode 6은 자체 답변 생성을 하므로 아래 공통 답변 생성 건너뜀
            else:
                context = text[:5000]
                print(f"[Mode 6] 관련 섹션을 찾지 못함. 문서 앞부분으로 대체합니다.")

        else:
            print("잘못된 입력입니다.")
            continue

        # --- 답변 생성 단계 ---
        print("\n[AI 답변 생성 중...]\n")
        prompt = f"""
당신은 유능한 AI 어시스턴트입니다. 아래의 [배경 지식]을 참고하여 사용자의 [질문]에 답변하세요.
만약 배경 지식에 정답이 없다면, 솔직하게 모른다고 말하세요.
반드시 사용자가 질문한 언어와 동일한 언어로 답변하세요 (예: 한국어 질문 -> 한국어 답변).

[배경 지식]
{context}

[질문]
{query}
"""
        try:
            client = ollama.Client(host=Config.OLLAMA_HOST)
            stream = client.chat(
                model=Config.MODEL_CHAT,
                messages=[{'role': 'user', 'content': prompt}],
                stream=True,
                options=get_llm_options()
            )
            for chunk in stream:
                print(chunk['message']['content'], end="", flush=True)
            print()
        except Exception as e:
            print(f"Ollama 오류: {e}")

if __name__ == "__main__":
    main()
