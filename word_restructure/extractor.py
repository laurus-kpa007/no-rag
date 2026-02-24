"""
문서 파싱/추출 모듈
python-docx를 사용하여 .docx 파일에서 모든 요소를 순서대로 추출합니다.
"""

import re
from dataclasses import dataclass, field
from typing import List, Optional
from docx import Document
from docx.oxml.ns import qn


@dataclass
class DocumentElement:
    """문서의 개별 요소를 표현하는 데이터 클래스"""
    type: str           # 'heading', 'paragraph', 'table', 'list_item'
    content: str        # 텍스트 내용
    level: int = 0      # heading level (1-9), list indent level
    style: str = ''     # 원본 Word 스타일 이름
    metadata: dict = field(default_factory=dict)  # 추가 정보


@dataclass
class ExtractedDocument:
    """추출된 문서 전체를 표현"""
    elements: List[DocumentElement]
    raw_text: str       # 전체 텍스트 (LLM 분석용)
    title: str = ''     # 문서 제목 (감지된 경우)


def get_heading_level(paragraph) -> int:
    """Paragraph의 Heading 레벨을 반환. Heading이 아니면 0 반환."""
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name.startswith('Heading'):
        try:
            return int(style_name.replace('Heading', '').strip())
        except ValueError:
            return 0
    # 한국어 스타일 처리
    if '제목' in style_name:
        for ch in style_name:
            if ch.isdigit():
                return int(ch)
        return 1
    # Title 스타일
    if style_name == 'Title':
        return 0  # 문서 제목 (특수 처리)
    return -1  # heading 아님


def is_list_style(paragraph) -> bool:
    """리스트 스타일인지 확인"""
    style_name = paragraph.style.name if paragraph.style else ""
    list_keywords = ['List', 'Bullet', 'Number', '목록', '글머리']
    return any(kw.lower() in style_name.lower() for kw in list_keywords)


def get_list_level(paragraph) -> int:
    """리스트의 들여쓰기 레벨을 반환"""
    # XML에서 ilvl (indent level) 추출
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            ilvl = numPr.find(qn('w:ilvl'))
            if ilvl is not None:
                try:
                    return int(ilvl.get(qn('w:val')))
                except (ValueError, TypeError):
                    pass
    return 0


def table_to_text(table) -> str:
    """Table 객체를 텍스트로 변환"""
    rows_text = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows_text.append(" | ".join(cells))
    return "\n".join(rows_text)


def table_to_data(table) -> dict:
    """Table 객체를 구조화된 데이터로 변환"""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)

    return {
        'rows': rows,
        'num_rows': len(rows),
        'num_cols': len(rows[0]) if rows else 0,
        'header': rows[0] if rows else [],
    }


def extract_document(file_path: str) -> ExtractedDocument:
    """
    .docx 파일에서 모든 요소를 순서대로 추출합니다.

    Args:
        file_path: .docx 파일 경로

    Returns:
        ExtractedDocument: 추출된 문서 객체
    """
    doc = Document(file_path)
    elements: List[DocumentElement] = []
    raw_text_parts: List[str] = []
    doc_title = ''

    # XML 레벨에서 body 요소 순회 (단락/표 순서 보장)
    body = doc.element.body

    # paragraph와 table 매핑 구성
    para_map = {para._element: para for para in doc.paragraphs}
    table_map = {table._element: table for table in doc.tables}

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            para = para_map.get(child)
            if para is None:
                continue

            text = para.text.strip()
            if not text:
                continue

            heading_level = get_heading_level(para)

            # 문서 제목 (Title 스타일)
            if heading_level == 0:
                doc_title = text
                elements.append(DocumentElement(
                    type='heading',
                    content=text,
                    level=0,
                    style=para.style.name if para.style else '',
                ))
                raw_text_parts.append(f"[제목] {text}")

            # Heading
            elif heading_level > 0:
                elements.append(DocumentElement(
                    type='heading',
                    content=text,
                    level=heading_level,
                    style=para.style.name if para.style else '',
                ))
                prefix = '#' * heading_level
                raw_text_parts.append(f"\n{prefix} {text}")

            # 리스트 아이템
            elif is_list_style(para):
                list_level = get_list_level(para)
                elements.append(DocumentElement(
                    type='list_item',
                    content=text,
                    level=list_level,
                    style=para.style.name if para.style else '',
                ))
                indent = '  ' * list_level
                raw_text_parts.append(f"{indent}- {text}")

            # 일반 단락
            else:
                elements.append(DocumentElement(
                    type='paragraph',
                    content=text,
                    level=0,
                    style=para.style.name if para.style else '',
                ))
                raw_text_parts.append(text)

        elif tag == 'tbl':
            table = table_map.get(child)
            if table is None:
                continue

            table_text = table_to_text(table)
            if table_text.strip():
                elements.append(DocumentElement(
                    type='table',
                    content=table_text,
                    level=0,
                    style='Table',
                    metadata=table_to_data(table),
                ))
                raw_text_parts.append(f"\n[표]\n{table_text}\n")

    raw_text = "\n".join(raw_text_parts)

    # 제목이 감지되지 않았으면 첫 heading이나 첫 줄을 사용
    if not doc_title:
        for elem in elements:
            if elem.type == 'heading':
                doc_title = elem.content
                break
        if not doc_title and elements:
            doc_title = elements[0].content[:50]

    return ExtractedDocument(
        elements=elements,
        raw_text=raw_text,
        title=doc_title,
    )


def elements_to_sections(elements: List[DocumentElement]) -> List[dict]:
    """
    요소 리스트를 섹션 단위로 그룹핑합니다.
    heading을 기준으로 섹션을 구분합니다.

    Returns:
        list of {'title': str, 'level': int, 'elements': List[DocumentElement], 'index': int}
    """
    sections = []
    current_section = {
        'title': '서두',
        'level': 0,
        'elements': [],
        'index': 0,
    }

    for elem in elements:
        if elem.type == 'heading' and elem.level > 0:
            # 이전 섹션 저장 (비어있지 않으면)
            if current_section['elements']:
                sections.append(current_section)

            current_section = {
                'title': elem.content,
                'level': elem.level,
                'elements': [elem],
                'index': len(sections),
            }
        else:
            current_section['elements'].append(elem)

    # 마지막 섹션 저장
    if current_section['elements']:
        sections.append(current_section)

    return sections


def sections_to_text_summary(sections: List[dict]) -> str:
    """섹션 리스트를 LLM 분석용 텍스트 요약으로 변환"""
    lines = []
    for i, section in enumerate(sections):
        level_prefix = '  ' * (section['level'] - 1) if section['level'] > 0 else ''
        content_preview = ''
        for elem in section['elements']:
            if elem.type != 'heading':
                content_preview += elem.content + ' '
        content_preview = content_preview.strip()[:200]
        lines.append(
            f"[{i}] {level_prefix}{section['title']} "
            f"(요소 {len(section['elements'])}개, 미리보기: {content_preview}...)"
        )
    return "\n".join(lines)
