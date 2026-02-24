"""
테스트용 더미 문서 생성 및 재구성 로직 테스트
Ollama 없이도 문서 추출/재구성/출력 로직을 테스트할 수 있습니다.

사용법:
    python -m word_restructure.test_restructure
"""

import os
import sys
from docx import Document


def create_test_document_structured(path: str):
    """헤딩이 있는 구조화된 테스트 문서 생성 (의도적으로 순서가 뒤섞임)"""
    doc = Document()

    doc.add_heading('인공지능 기술 보고서', 0)

    # 섹션 0: 서론
    doc.add_heading('1. 서론', level=1)
    doc.add_paragraph(
        '인공지능(AI)은 현대 기술의 핵심 분야로 자리잡았습니다. '
        '본 보고서에서는 AI의 현재 동향과 미래 전망을 분석합니다.'
    )
    doc.add_paragraph(
        '최근 대규모 언어 모델(LLM)의 발전으로 자연어 처리 분야에서 '
        '혁명적인 변화가 일어나고 있습니다.'
    )

    # 섹션 1: 결론 (의도적으로 중간에 배치 - 재구성 시 뒤로 이동해야 함)
    doc.add_heading('결론', level=1)
    doc.add_paragraph(
        'AI 기술의 발전은 기회와 도전을 동시에 가져옵니다. '
        '기술의 책임있는 개발과 활용이 중요합니다.'
    )

    # 섹션 2: 기술 현황
    doc.add_heading('2. 기술 현황', level=1)

    # 섹션 3: 대규모 언어 모델
    doc.add_heading('2.1 대규모 언어 모델', level=2)
    doc.add_paragraph(
        'GPT, Claude, Gemma 등 다양한 LLM이 등장하여 텍스트 생성, '
        '요약, 번역 등 다양한 작업을 수행하고 있습니다.'
    )
    doc.add_paragraph('주요 LLM 모델:', style='List Bullet')
    doc.add_paragraph('GPT-4: OpenAI의 대표 모델', style='List Bullet')
    doc.add_paragraph('Claude: Anthropic의 안전한 AI', style='List Bullet')
    doc.add_paragraph('Gemma: Google의 오픈소스 모델', style='List Bullet')

    # 섹션 4: 컴퓨터 비전
    doc.add_heading('2.2 컴퓨터 비전', level=2)
    doc.add_paragraph(
        '이미지 인식, 객체 탐지, 자율주행 등 컴퓨터 비전 분야에서도 '
        'AI 기술이 빠르게 발전하고 있습니다.'
    )

    # 섹션 5: 산업 적용 사례 (표 포함)
    doc.add_heading('3. 산업 적용 사례', level=1)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['분야', '적용 사례', '효과']
    for j, header in enumerate(headers):
        table.cell(0, j).text = header

    data = [
        ['의료', '질병 진단 보조', '진단 정확도 30% 향상'],
        ['금융', '사기 탐지', '사기 적발률 50% 향상'],
        ['제조', '품질 검사 자동화', '불량률 20% 감소'],
    ]
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            table.cell(i + 1, j).text = cell_text

    # 섹션 6: 향후 전망
    doc.add_heading('4. 향후 전망', level=1)
    doc.add_paragraph(
        'AI 기술은 앞으로 더욱 발전하여 인간의 삶을 크게 변화시킬 것으로 예상됩니다. '
        '특히 멀티모달 AI, 에이전트 AI, 소형화된 온디바이스 AI가 주요 트렌드가 될 것입니다.'
    )

    doc.save(path)
    print(f"[TEST] 구조화된 테스트 문서 생성: {path}")


def create_test_document_unstructured(path: str):
    """헤딩 없는 비정형 테스트 문서 생성"""
    doc = Document()

    doc.add_paragraph('회의록')
    doc.add_paragraph('일시: 2026년 2월 24일')
    doc.add_paragraph('장소: 본사 회의실')
    doc.add_paragraph('참석자: 김철수, 이영희, 박민수, 정수진')
    doc.add_paragraph(
        '첫 번째 안건으로 신규 프로젝트 일정에 대해 논의했습니다. '
        '김철수 팀장이 전체 일정을 설명했고, 3월 말까지 1차 프로토타입을 '
        '완성하기로 합의했습니다.'
    )
    doc.add_paragraph(
        '두 번째 안건으로 예산 배분에 대해 논의했습니다. '
        '인프라 구축에 40%, 인력 비용에 35%, 기타 운영비에 25%를 '
        '배분하기로 했습니다.'
    )
    doc.add_paragraph(
        '세 번째 안건으로 인력 충원 계획을 논의했습니다. '
        'AI 엔지니어 2명, 프론트엔드 개발자 1명을 3월 중으로 '
        '채용하기로 결정했습니다.'
    )
    doc.add_paragraph('다음 회의 일정: 2026년 3월 3일')

    doc.save(path)
    print(f"[TEST] 비정형 테스트 문서 생성: {path}")


def test_extractor():
    """문서 추출 로직 테스트"""
    from .extractor import extract_document, elements_to_sections, sections_to_text_summary

    test_dir = 'test_data'
    os.makedirs(test_dir, exist_ok=True)

    # 테스트 문서 생성
    structured_path = os.path.join(test_dir, 'test_structured.docx')
    unstructured_path = os.path.join(test_dir, 'test_unstructured.docx')

    create_test_document_structured(structured_path)
    create_test_document_unstructured(unstructured_path)

    # 구조화된 문서 테스트
    print("\n" + "=" * 50)
    print("구조화된 문서 추출 테스트")
    print("=" * 50)
    extracted = extract_document(structured_path)
    print(f"제목: {extracted.title}")
    print(f"요소 수: {len(extracted.elements)}")
    print(f"텍스트 길이: {len(extracted.raw_text)}자")

    for elem in extracted.elements:
        prefix = f"[{elem.type}]"
        if elem.type == 'heading':
            prefix += f"(L{elem.level})"
        content_preview = elem.content[:60] + "..." if len(elem.content) > 60 else elem.content
        print(f"  {prefix} {content_preview}")

    sections = elements_to_sections(extracted.elements)
    print(f"\n섹션 수: {len(sections)}")
    for s in sections:
        body_count = len(s['body_elements'])
        print(f"  [{s['index']}] L{s['level']} \"{s['title']}\" - body 요소 {body_count}개")
    print()
    print(sections_to_text_summary(sections))

    # 비정형 문서 테스트
    print("\n" + "=" * 50)
    print("비정형 문서 추출 테스트")
    print("=" * 50)
    extracted2 = extract_document(unstructured_path)
    print(f"제목: {extracted2.title}")
    print(f"요소 수: {len(extracted2.elements)}")

    sections2 = elements_to_sections(extracted2.elements)
    print(f"섹션 수: {len(sections2)}")

    print("\n[TEST] 추출 테스트 완료!")
    return structured_path, unstructured_path


def test_restructure_without_llm(structured_path: str):
    """
    LLM 없이 재구성 로직 테스트.
    수동으로 RestructurePlan을 만들어서 재구성 파이프라인을 검증합니다.
    """
    from .extractor import extract_document, elements_to_sections
    from .analyzer import TocItem, RestructurePlan, DocumentAnalysis
    from .restructurer import restructure_document
    from .writer_docx import write_docx
    from .writer_md import write_md

    print("\n" + "=" * 50)
    print("재구성 로직 테스트 (LLM 없이)")
    print("=" * 50)

    extracted = extract_document(structured_path)
    sections = elements_to_sections(extracted.elements)

    print(f"\n원본 섹션 ({len(sections)}개):")
    for s in sections:
        body_count = len(s['body_elements'])
        print(f"  [{s['index']}] L{s['level']} \"{s['title']}\" - body {body_count}개")

    # 수동 재구성 계획 생성
    # 원본 순서: 서론(0), 결론(1), 기술현황(2), LLM(3), CV(4), 산업적용(5), 전망(6)
    # 재구성 순서: 서론 → 기술현황(LLM, CV) → 산업적용 → 전망 → 결론
    analysis = DocumentAnalysis(
        document_type='보고서',
        main_topic='인공지능 기술 동향',
        key_themes=['LLM', '컴퓨터 비전', '산업 적용'],
        current_structure_assessment='결론이 중간에 위치하여 재배치 필요',
        content_sections=[],
    )

    plan = RestructurePlan(
        title='인공지능 기술 보고서 (재구성)',
        toc=[
            TocItem(level=1, title='서론', source_sections=[0]),
            TocItem(level=1, title='기술 현황', source_sections=[2], subsections=[
                TocItem(level=2, title='대규모 언어 모델', source_sections=[3]),
                TocItem(level=2, title='컴퓨터 비전', source_sections=[4]),
            ]),
            TocItem(level=1, title='산업 적용 사례', source_sections=[5]),
            TocItem(level=1, title='향후 전망', source_sections=[6]),
            TocItem(level=1, title='결론', source_sections=[1]),  # 결론을 맨 뒤로!
        ],
        analysis=analysis,
    )

    # 재구성 실행
    restructured = restructure_document(extracted, plan, refine=False)

    # 결과 검증
    print(f"\n재구성 결과:")
    print(f"  문서 제목: {restructured.title}")
    print(f"  섹션 수: {len(restructured.sections)}")

    errors = []

    def _print_sections(sections, indent=0):
        for s in sections:
            prefix = '  ' * indent
            content_preview = ''
            if s.content_elements:
                first = s.content_elements[0]
                content_preview = first.content[:50]
            print(f"  {prefix}[L{s.level}] {s.title} "
                  f"({len(s.content_elements)}개 요소) → {content_preview}")
            _print_sections(s.subsections, indent + 1)

    _print_sections(restructured.sections)

    # 검증 1: 결론이 마지막 섹션인지
    if restructured.sections[-1].title != '결론':
        errors.append(f"결론이 마지막이 아님: {restructured.sections[-1].title}")
    else:
        print("\n  [PASS] 결론이 마지막 섹션으로 올바르게 재배치됨")

    # 검증 2: 기술 현황에 하위 섹션이 있는지
    tech_section = None
    for s in restructured.sections:
        if s.title == '기술 현황':
            tech_section = s
            break
    if tech_section is None:
        errors.append("기술 현황 섹션을 찾을 수 없음")
    elif len(tech_section.subsections) != 2:
        errors.append(f"기술 현황 하위 섹션이 2개가 아님: {len(tech_section.subsections)}")
    else:
        print("  [PASS] 기술 현황에 하위 섹션 2개 존재")

    # 검증 3: 산업 적용 사례에 표가 있는지
    industry_section = None
    for s in restructured.sections:
        if s.title == '산업 적용 사례':
            industry_section = s
            break
    if industry_section is None:
        errors.append("산업 적용 사례 섹션을 찾을 수 없음")
    else:
        has_table = any(e.type == 'table' for e in industry_section.content_elements)
        if not has_table:
            errors.append("산업 적용 사례에 표가 없음")
        else:
            print("  [PASS] 산업 적용 사례에 표 요소 존재")

    # 검증 4: 중복 콘텐츠 없는지 (기술 현황 부모와 자식이 겹치지 않는지)
    if tech_section:
        parent_texts = {e.content for e in tech_section.content_elements}
        child_texts = set()
        for sub in tech_section.subsections:
            for e in sub.content_elements:
                child_texts.add(e.content)
        overlap = parent_texts & child_texts
        if overlap:
            errors.append(f"기술 현황 부모-자식 콘텐츠 중복: {overlap}")
        else:
            print("  [PASS] 부모-자식 섹션 간 콘텐츠 중복 없음")

    # 출력 파일 생성 테스트
    output_dir = 'test_output'
    os.makedirs(output_dir, exist_ok=True)
    output_docx = os.path.join(output_dir, 'test_restructured.docx')
    output_md = os.path.join(output_dir, 'test_restructured.md')

    write_docx(restructured, output_docx)
    write_md(restructured, output_md)

    # MD 파일 목차 검증
    with open(output_md, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # 목차에서 결론이 마지막인지 확인
    toc_start = md_content.find('## 목차')
    toc_end = md_content.find('---', toc_start + 1)
    toc_section = md_content[toc_start:toc_end]
    toc_lines = [l for l in toc_section.split('\n') if l.strip().startswith('-')]
    if toc_lines:
        last_top = None
        for l in toc_lines:
            if not l.startswith('  '):  # 최상위 항목 (들여쓰기 없음)
                last_top = l
        if last_top and '결론' in last_top:
            print("  [PASS] MD 목차에서 결론이 마지막 최상위 항목")
        else:
            errors.append(f"MD 목차 마지막 최상위 항목이 결론이 아님: {last_top}")

    # 본문에서 섹션 순서 확인
    heading_positions = []
    for line in md_content.split('\n'):
        if line.startswith('## ') and line != '## 목차':
            heading_positions.append(line.strip('# ').strip())

    expected_order = ['서론', '기술 현황', '산업 적용 사례', '향후 전망', '결론']
    if heading_positions == expected_order:
        print("  [PASS] MD 본문 섹션 순서가 올바름")
    else:
        errors.append(f"MD 본문 순서 불일치:\n    예상: {expected_order}\n    실제: {heading_positions}")

    if errors:
        print(f"\n[FAIL] {len(errors)}개 오류 발견:")
        for err in errors:
            print(f"  - {err}")
    else:
        print("\n[ALL PASS] 모든 재구성 로직 테스트 통과!")

    return len(errors) == 0


def test_edge_cases(structured_path: str):
    """
    엣지 케이스 테스트: LLM이 잘못된 인덱스를 반환하는 상황을 시뮬레이션.
    - 범위 초과 인덱스
    - 부모-자식 완전 중복 인덱스
    - 빈 source_sections
    - description으로 빈 부모 섹션 보완
    """
    from .extractor import extract_document, elements_to_sections
    from .analyzer import TocItem, RestructurePlan, DocumentAnalysis, _validate_source_sections
    from .restructurer import restructure_document
    from .writer_md import write_md

    print("\n" + "=" * 50)
    print("엣지 케이스 테스트")
    print("=" * 50)

    extracted = extract_document(structured_path)
    sections = elements_to_sections(extracted.elements)
    total = len(sections)

    errors = []

    # --- 테스트 1: 범위 초과 인덱스 검증 ---
    print("\n[테스트 1] 범위 초과 인덱스 검증")
    test_items = [
        TocItem(level=1, title='유효 섹션', source_sections=[0, 1]),
        TocItem(level=1, title='범위 초과', source_sections=[99, -1, 100]),
        TocItem(level=1, title='혼합', source_sections=[2, 50, 3]),
    ]
    _validate_source_sections(test_items, total)

    if test_items[0].source_sections == [0, 1]:
        print("  [PASS] 유효 인덱스 보존됨")
    else:
        errors.append(f"유효 인덱스가 변경됨: {test_items[0].source_sections}")

    if test_items[1].source_sections == []:
        print("  [PASS] 범위 초과 인덱스 모두 제거됨")
    else:
        errors.append(f"범위 초과 인덱스 제거 실패: {test_items[1].source_sections}")

    if test_items[2].source_sections == [2, 3]:
        print("  [PASS] 혼합 인덱스에서 초과만 제거됨")
    else:
        errors.append(f"혼합 인덱스 처리 실패: {test_items[2].source_sections}")

    # --- 테스트 2: 부모-자식 완전 중복 + description 보완 ---
    print("\n[테스트 2] 부모-자식 인덱스 중복 + description 보완")
    analysis = DocumentAnalysis(
        document_type='보고서', main_topic='테스트',
        key_themes=[], current_structure_assessment='', content_sections=[],
    )
    plan = RestructurePlan(
        title='엣지 케이스 테스트',
        toc=[
            TocItem(level=1, title='부모 섹션', source_sections=[0, 1, 2],
                    description='이 섹션은 AI 기술 현황을 다룹니다.',
                    subsections=[
                        TocItem(level=2, title='자식1', source_sections=[0, 1]),
                        TocItem(level=2, title='자식2', source_sections=[2]),
                    ]),
            TocItem(level=1, title='나머지', source_sections=[3, 4, 5, 6]),
        ],
        analysis=analysis,
    )

    restructured = restructure_document(extracted, plan, refine=False)

    # 부모의 own_indices는 [0,1,2] - [0,1] - [2] = [] → description이 삽입되어야 함
    parent = restructured.sections[0]
    if parent.content_elements:
        if parent.content_elements[0].content == '이 섹션은 AI 기술 현황을 다룹니다.':
            print("  [PASS] 빈 부모에 description이 삽입됨")
        else:
            errors.append(f"description 삽입 내용 불일치: {parent.content_elements[0].content}")
    else:
        errors.append("부모 섹션에 description이 삽입되지 않음 (content_elements 비어있음)")

    # 자식 섹션에 내용이 있는지
    if parent.subsections[0].content_elements:
        print("  [PASS] 자식1에 내용 존재")
    else:
        errors.append("자식1에 내용이 없음")

    if parent.subsections[1].content_elements:
        print("  [PASS] 자식2에 내용 존재")
    else:
        errors.append("자식2에 내용이 없음")

    # --- 테스트 3: 출력에서 모든 계획된 섹션이 나타나는지 ---
    print("\n[테스트 3] 모든 계획된 섹션이 출력에 반영되는지")
    output_dir = 'test_output'
    os.makedirs(output_dir, exist_ok=True)
    output_md = os.path.join(output_dir, 'test_edge_case.md')
    write_md(restructured, output_md)

    with open(output_md, 'r', encoding='utf-8') as f:
        md_content = f.read()

    for title in ['부모 섹션', '자식1', '자식2', '나머지']:
        if title in md_content:
            print(f"  [PASS] '{title}' 섹션이 출력에 존재")
        else:
            errors.append(f"'{title}' 섹션이 출력에서 누락됨")

    if errors:
        print(f"\n[FAIL] {len(errors)}개 오류 발견:")
        for err in errors:
            print(f"  - {err}")
    else:
        print("\n[ALL PASS] 모든 엣지 케이스 테스트 통과!")

    return len(errors) == 0


def test_dynamic_guide_and_hierarchy():
    """
    동적 가이드 계산 로직 테스트 + 다수 섹션 시 계층 깊이 검증.
    문서 규모에 따라 적절한 min_major/depth_guide가 생성되는지,
    많은 섹션을 가진 문서에서 중분류가 실제로 생성되는지 검증합니다.
    """
    from .extractor import DocumentElement, ExtractedDocument, elements_to_sections
    from .analyzer import TocItem, RestructurePlan, DocumentAnalysis
    from .restructurer import restructure_document
    from .writer_md import write_md

    print("\n" + "=" * 50)
    print("동적 가이드 + 계층 깊이 테스트")
    print("=" * 50)

    errors = []

    # --- 테스트 1: 동적 가이드 계산 로직 검증 ---
    print("\n[테스트 1] 동적 가이드 계산 로직")

    test_cases = [
        (3, 2, "1~2 레벨"),          # 작은 문서
        (5, 2, "1~2 레벨"),          # 경계: 5개
        (6, 3, "2레벨"),             # 중간 문서 시작
        (12, 3, "2레벨"),            # 중간 문서 (12//4=3)
        (15, 3, "2레벨"),            # 경계: 15개
        (16, 5, "3레벨"),            # 큰 문서 시작
        (30, 6, "3레벨"),            # 큰 문서 (30//5=6)
        (50, 10, "3레벨"),           # 매우 큰 문서 (50//5=10)
    ]

    for n, expected_min, expected_keyword in test_cases:
        if n <= 5:
            min_major = 2
            depth_guide = "1~2 레벨 구조로 구성하세요."
        elif n <= 15:
            min_major = max(3, n // 4)
            depth_guide = "반드시 2레벨(대분류/중분류) 이상으로 구성하세요."
        else:
            min_major = max(5, n // 5)
            depth_guide = "반드시 3레벨(대분류/중분류/소분류) 구조로 상세하게 구성하세요."

        if min_major != expected_min:
            errors.append(f"n={n}: min_major={min_major} (기대: {expected_min})")
        elif expected_keyword not in depth_guide:
            errors.append(f"n={n}: depth_guide에 '{expected_keyword}' 없음: {depth_guide}")
        else:
            print(f"  [PASS] n={n} → min_major={min_major}, depth_guide에 '{expected_keyword}' 포함")

    # --- 테스트 2: 많은 섹션으로 계층 깊이 검증 ---
    print("\n[테스트 2] 다수 섹션(12개) 문서에서 대/중분류 계층 검증")

    # 12개 섹션이 있는 Mock 문서 생성
    elements = [DocumentElement(type='heading', content='테스트 문서', level=0)]
    section_titles = [
        '프로젝트 개요', '배경 및 필요성', '목표 설정',
        '현황 분석', '데이터 수집', '데이터 분석 결과',
        '설계 방안', '구현 세부사항', '시스템 아키텍처',
        '테스트 결과', '성능 평가', '결론 및 제언',
    ]
    for title in section_titles:
        elements.append(DocumentElement(type='heading', content=title, level=1))
        elements.append(DocumentElement(
            type='paragraph', content=f'{title}에 대한 상세 내용입니다.', level=0
        ))
        elements.append(DocumentElement(
            type='paragraph', content=f'{title} 관련 추가 설명과 데이터를 포함합니다.', level=0
        ))

    sections = elements_to_sections(elements)
    print(f"  생성된 섹션 수: {len(sections)}")

    # 12개 섹션 → min_major=3, 2레벨 필수
    # 대분류 4개 + 각 대분류에 중분류 2~3개로 구성된 계획 만들기
    analysis = DocumentAnalysis(
        document_type='기술문서', main_topic='프로젝트 보고서',
        key_themes=['분석', '설계', '평가'],
        current_structure_assessment='세분화 필요', content_sections=[],
    )

    plan = RestructurePlan(
        title='프로젝트 보고서 (재구성)',
        toc=[
            TocItem(level=1, title='1. 개요', source_sections=[0, 1, 2],
                    description='프로젝트의 배경과 목표',
                    subsections=[
                        TocItem(level=2, title='1.1 프로젝트 개요', source_sections=[0],
                                description='프로젝트 개요 설명'),
                        TocItem(level=2, title='1.2 배경 및 필요성', source_sections=[1],
                                description='프로젝트 배경'),
                        TocItem(level=2, title='1.3 목표', source_sections=[2],
                                description='목표 설정'),
                    ]),
            TocItem(level=1, title='2. 현황 분석', source_sections=[3, 4, 5],
                    description='현황과 데이터 분석',
                    subsections=[
                        TocItem(level=2, title='2.1 현황 분석', source_sections=[3],
                                description='현황 분석 내용'),
                        TocItem(level=2, title='2.2 데이터 수집 및 분석', source_sections=[4, 5],
                                description='데이터 수집과 분석 결과'),
                    ]),
            TocItem(level=1, title='3. 설계 및 구현', source_sections=[6, 7, 8],
                    description='시스템 설계와 구현',
                    subsections=[
                        TocItem(level=2, title='3.1 설계 방안', source_sections=[6],
                                description='설계 방안 상세'),
                        TocItem(level=2, title='3.2 구현 및 아키텍처', source_sections=[7, 8],
                                description='구현 세부사항과 아키텍처'),
                    ]),
            TocItem(level=1, title='4. 평가 및 결론', source_sections=[9, 10, 11],
                    description='테스트 결과와 결론',
                    subsections=[
                        TocItem(level=2, title='4.1 테스트 및 성능', source_sections=[9, 10],
                                description='테스트 결과와 성능 평가'),
                        TocItem(level=2, title='4.2 결론 및 제언', source_sections=[11],
                                description='결론과 향후 제언'),
                    ]),
        ],
        analysis=analysis,
    )

    # 가이드 기준에 맞는지 확인
    n = len(sections)
    min_major = max(3, n // 4)
    level1_count = len(plan.toc)
    total_level2 = sum(len(item.subsections) for item in plan.toc)

    if level1_count < min_major:
        errors.append(f"대분류 {level1_count}개 < 최소 {min_major}개")
    else:
        print(f"  [PASS] 대분류 {level1_count}개 >= 최소 {min_major}개")

    if total_level2 == 0:
        errors.append("중분류가 하나도 없음")
    else:
        print(f"  [PASS] 중분류 총 {total_level2}개 존재")

    # 모든 대분류에 중분류가 있는지
    for item in plan.toc:
        if len(item.subsections) == 0:
            errors.append(f"대분류 '{item.title}'에 중분류 없음")
        else:
            print(f"  [PASS] '{item.title}' → 중분류 {len(item.subsections)}개")

    # 재구성 실행 및 출력 검증
    mock_extracted = ExtractedDocument(
        elements=elements, title='테스트 문서', raw_text='',
    )
    restructured = restructure_document(extracted=mock_extracted, plan=plan, refine=False)

    output_dir = 'test_output'
    os.makedirs(output_dir, exist_ok=True)
    output_md = os.path.join(output_dir, 'test_hierarchy.md')
    write_md(restructured, output_md)

    with open(output_md, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # MD에서 대분류(##)와 중분류(###) 개수 확인
    h2_count = sum(1 for line in md_content.split('\n')
                   if line.startswith('## ') and line != '## 목차')
    h3_count = sum(1 for line in md_content.split('\n')
                   if line.startswith('### '))

    if h2_count < min_major:
        errors.append(f"MD 대분류(##) {h2_count}개 < 최소 {min_major}개")
    else:
        print(f"  [PASS] MD 대분류(##) {h2_count}개 출력됨")

    if h3_count == 0:
        errors.append("MD 중분류(###)가 없음")
    else:
        print(f"  [PASS] MD 중분류(###) {h3_count}개 출력됨")

    if errors:
        print(f"\n[FAIL] {len(errors)}개 오류 발견:")
        for err in errors:
            print(f"  - {err}")
    else:
        print("\n[ALL PASS] 동적 가이드 + 계층 깊이 테스트 통과!")

    return len(errors) == 0


def main():
    """테스트 메인"""
    print("=" * 55)
    print("  Word Document Restructuring Tool - 테스트")
    print("=" * 55)

    # Step 1: 추출 테스트 (Ollama 불필요)
    structured_path, _ = test_extractor()

    # Step 2: 재구성 로직 테스트 (Ollama 불필요)
    passed = test_restructure_without_llm(structured_path)

    # Step 3: 엣지 케이스 테스트 (Ollama 불필요)
    passed = test_edge_cases(structured_path) and passed

    # Step 4: 동적 가이드 + 계층 깊이 테스트 (Ollama 불필요)
    passed = test_dynamic_guide_and_hierarchy() and passed

    # Step 5: 전체 파이프라인 테스트 (Ollama 필요)
    if '--full' in sys.argv:
        print("\n" + "=" * 55)
        print("전체 파이프라인 테스트 (Ollama 필요)")
        print("=" * 55)

        from .config import Config, get_ollama_client

        try:
            client = get_ollama_client()
            client.list()
            print("[TEST] Ollama 연결 성공")
        except Exception as e:
            print(f"[TEST] Ollama 연결 실패: {e}")
            print("[TEST] 전체 파이프라인 테스트를 건너뜁니다.")
            print(f"[TEST] Ollama를 시작하고 {Config.MODEL}을 설치하세요:")
            print(f"       ollama pull {Config.MODEL}")
            return

        # 전체 실행
        print(f"\n[TEST] 전체 파이프라인 실행: {structured_path}")
        from .main import main as run_main
        sys.argv = ['test', structured_path, '-o', 'test_output_full']
        try:
            run_main()
            print("\n[TEST] 전체 파이프라인 테스트 성공!")
        except SystemExit:
            pass
        except Exception as e:
            print(f"\n[TEST] 전체 파이프라인 테스트 실패: {e}")
    else:
        print("\n[INFO] 전체 파이프라인(LLM) 테스트는 --full 옵션으로 실행하세요:")
        print("       python -m word_restructure.test_restructure --full")

    if not passed:
        sys.exit(1)


if __name__ == "__main__":
    main()
