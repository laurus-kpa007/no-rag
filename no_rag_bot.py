# 필요한 라이브러리 설치 명령어:
# pip install ollama python-docx

import os
import sys
import ollama
from docx import Document

# 설정
DOC_PATH = 'data.docx'            # 분석할 파일 경로 (.docx 또는 .md)
MODEL_NAME = 'gemma3:4b'           # 사용할 Ollama 모델 이름
OLLAMA_HOST = 'http://localhost:11434' # Ollama 서버 주소 (기본값: http://localhost:11434)
NUM_CTX = 32768                   # 컨텍스트 윈도우 크기 (32K = 32768). 문서 크기에 맞춰 조절하세요.

def extract_text_from_docx(file_path):
    """
    .docx 파일에서 모든 텍스트를 추출합니다.
    """
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
        return "\n".join(full_text)
    except Exception as e:
        print(f"docx 읽기 오류: {e}")
        return None

def extract_text_from_md(file_path):
    """
    .md 또는 텍스트 파일에서 내용을 추출합니다.
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"파일 읽기 오류: {e}")
        return None

def load_document(file_path):
    """
    확장자에 따라 적절한 텍스트 추출 함수를 호출합니다.
    """
    if not os.path.exists(file_path):
        print(f"오류: '{file_path}' 파일을 찾을 수 없습니다.")
        return None

    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext in ['.md', '.txt']:
        return extract_text_from_md(file_path)
    else:
        print(f"지원하지 않는 파일 형식입니다: {ext}")
        return None

def chat_with_doc(doc_content):
    """
    문서 내용을 컨텍스트로 사용하여 사용자와 대화합니다.
    """
    print(f"\n[INFO] 문서 내용이 로드되었습니다. ({len(doc_content)} 자)")
    print("[INFO] 질의응답을 시작합니다. (종료하려면 'q' 또는 'exit' 입력)")
    print("-" * 50)

    while True:
        try:
            user_input = input("\n질문: ").strip()
        except KeyboardInterrupt:
            print("\n프로그램을 종료합니다.")
            break

        if user_input.lower() in ['q', 'exit']:
            print("프로그램을 종료합니다.")
            break

        if not user_input:
            continue

        # 프롬프트 구성 (Full Context Stuffing)
        prompt = f"""
당신은 도움이 되는 AI 어시스턴트입니다. 아래의 문서 내용을 바탕으로 사용자의 질문에 답변해 주세요.
문서 내용에서 정답을 찾을 수 없다면, 모른다고 답변해 주세요.
반드시 사용자가 질문한 언어와 동일한 언어로 답변하세요 (예: 한국어 질문 -> 한국어 답변, 영어 질문 -> 영어 답변).

[문서 내용]
{doc_content}

[사용자 질문]
{user_input}
"""

        print("\n답변: ", end="", flush=True)

        try:
            # Ollama 클라이언트 생성 (호스트 지정)
            client = ollama.Client(host=OLLAMA_HOST)

            # 스트리밍 호출
            stream = client.chat(
                model=MODEL_NAME,
                messages=[{'role': 'user', 'content': prompt}],
                stream=True,
                options={'num_ctx': NUM_CTX} # 컨텍스트 크기 설정
            )

            for chunk in stream:
                content = chunk['message']['content']
                print(content, end="", flush=True)
            
            print() # 줄바꿈

        except Exception as e:
            print(f"\n[오류] Ollama 통신 중 문제 발생: {e}")
            print("Ollama가 실행 중인지, 모델이 설치되어 있는지 확인해주세요.")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        DOC_PATH = sys.argv[1]

    print(f"'{DOC_PATH}' 문서를 읽는 중...")
    document_content = load_document(DOC_PATH)

    if document_content:
        chat_with_doc(document_content)
